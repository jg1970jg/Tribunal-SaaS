"""
MAIN - LexForum (FastAPI)
============================================================
Servidor principal com:
  - Conexão ao Supabase
  - Rota de saúde (GET /health)
  - Rota protegida de teste (GET /me)
  - Rota de análise (POST /analyze) com suporte a tiers
  - Endpoints de wallet e admin
============================================================
"""

import asyncio
import collections
import itertools
import io
import os
import logging
import re
import secrets
import signal
import sys
from contextlib import asynccontextmanager
from datetime import datetime, timezone, timedelta
from typing import Any, Optional

from dotenv import load_dotenv
from fastapi import Depends, FastAPI, Form, Query, Request, UploadFile, File, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field, field_validator
from slowapi import Limiter
from slowapi.errors import RateLimitExceeded
from slowapi.util import get_remote_address
import threading

from auth_service import get_current_user, get_supabase, get_supabase_admin
from src.engine import (
    executar_analise_documento,
    InsufficientBalanceError,
    InvalidDocumentError,
    MissingApiKeyError,
    EngineError,
    get_wallet_manager,
)
from src.cost_controller import BudgetExceededError

# =============================================================================
# IN-MEMORY LOG BUFFER — para endpoint /admin/logs (monitorização remota)
# =============================================================================

class InMemoryLogHandler(logging.Handler):
    """Circular buffer que guarda os últimos N log records em memória."""

    def __init__(self, capacity: int = 2000):
        super().__init__()
        self._buffer: collections.deque = collections.deque(maxlen=capacity)
        self._counter = itertools.count(1)
        self._lock = threading.Lock()

    def emit(self, record: logging.LogRecord):
        try:
            seq = next(self._counter)
            with self._lock:
                self._buffer.append({
                    "id": seq,
                    "ts": datetime.fromtimestamp(record.created, tz=timezone.utc).isoformat(),
                    "level": record.levelname,
                    "logger": record.name,
                    "msg": self.format(record),
                })
        except Exception:
            pass  # Never break the app for logging

    def get_logs(self, limit: int = 200, level: str = None, search: str = None, since_id: int = 0):
        """Return filtered logs from buffer."""
        with self._lock:
            snapshot = list(self._buffer)
        result = []
        for entry in snapshot:
            if entry["id"] <= since_id:
                continue
            if level and entry["level"] != level.upper():
                continue
            if search and search.lower() not in entry["msg"].lower():
                continue
            result.append(entry)
        return result[-limit:]


# Install handler on root logger to capture ALL logs
_log_buffer = InMemoryLogHandler(capacity=5000)
_log_buffer.setFormatter(logging.Formatter("%(name)s | %(message)s"))
logging.root.addHandler(_log_buffer)
logging.root.setLevel(logging.INFO)

logger = logging.getLogger(__name__)

# Constantes de validação de ficheiros (usadas em /analyze e /analyze/add)
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".doc"}
MAX_QUESTION_LENGTH = 5000
MAX_ADDITIONAL_DOCS = 10
ADMIN_SESSION_TTL = 3600  # 1 hora
VALID_AREAS = {"Civil", "Penal", "Trabalho", "Administrativo", "Fiscal", "Comercial", "Família", "Outro"}


def _sanitize_filename(filename: str) -> str:
    """Sanitize filename to prevent path traversal and prompt injection."""
    if not filename:
        return "documento"
    # Remove path separators and keep only safe chars
    name = os.path.basename(filename)
    name = re.sub(r'[^\w._\-\s]', '_', name)
    return name[:255] if name else "documento"


# FIX 2026-02-14: Carregar .env ANTES de ler ADMIN_EMAILS
load_dotenv()

# Lista de emails de admin (para endpoints de admin)
ADMIN_EMAILS = [e.strip().lower() for e in os.environ.get("ADMIN_EMAILS", "").split(",") if e.strip()]

# Admin session tokens (from /admin/verify) — with 1h expiry
_admin_sessions: dict = {}
_admin_sessions_lock = threading.Lock()


# ============================================================
# RATE LIMITING
# ============================================================

def _get_user_or_ip(request: Request) -> str:
    """Usa IP do request para rate limiting (seguro, sem depender de JWT)."""
    return get_remote_address(request)


limiter = Limiter(key_func=_get_user_or_ip)


def _rate_limit_handler(request: Request, exc: RateLimitExceeded) -> JSONResponse:
    limit_detail = str(exc.detail or "")
    is_daily_block = "per 1 day" in limit_detail or "per day" in limit_detail

    if is_daily_block:
        _register_security_alert(request, limit_detail)

    return JSONResponse(
        status_code=429,
        content={
            "detail": "Demasiados pedidos. Tente novamente em breve.",
            "retry_after": limit_detail,
        },
    )


def _extract_user_info(request: Request) -> dict:
    """Extrai IP do request para logging de segurança."""
    ip = get_remote_address(request)
    return {"ip": ip, "user_id": None, "email": None}


def _register_security_alert(request: Request, limit_detail: str):
    """Regista alerta de segurança quando limite diário é atingido."""
    user_info = _extract_user_info(request)
    endpoint = request.url.path

    # Construir descrição legível
    parts = [f"Limite diário atingido: {limit_detail}"]
    parts.append(f"Endpoint: {endpoint}")
    parts.append(f"IP: {user_info['ip']}")
    if user_info["email"]:
        parts.append(f"Email: {user_info['email']}")
    if user_info["user_id"]:
        parts.append(f"User ID: {user_info['user_id']}")
    detail = " | ".join(parts)

    # Identificar o ofensor (email > user_id > IP)
    offender = user_info["email"] or user_info["user_id"] or user_info["ip"]

    logger.critical(
        f"[SECURITY ALERT] Bloqueio 24h activado! "
        f"Ofensor: {offender} | IP: {user_info['ip']} | Endpoint: {endpoint}"
    )

    try:
        from auth_service import get_supabase_admin
        sb = get_supabase_admin()
        sb.table("security_alerts").insert({
            "alert_type": "daily_rate_limit",
            "endpoint": endpoint,
            "offender": offender,
            "detail": detail,
        }).execute()
    except Exception as e:
        logger.error(f"Erro ao registar alerta de segurança: {e}")


# ============================================================
# BLACKLIST - Bloqueio de emails, IPs e domínios
# ============================================================

# Cache local (recarregado a cada 60s para não consultar DB a cada request)
_blacklist_cache: dict = {"emails": set(), "ips": set(), "domains": set(), "loaded_at": 0}
_blacklist_lock = threading.Lock()  # Thread-safe access to _blacklist_cache
BLACKLIST_CACHE_TTL = 60  # Recarregar a cada 60 segundos


def _load_blacklist():
    """Carrega blacklist do Supabase para cache local."""
    import time as _time
    now = _time.time()
    if now - _blacklist_cache["loaded_at"] < BLACKLIST_CACHE_TTL:
        return  # Cache ainda válido

    with _blacklist_lock:
        # Re-read timestamp inside the lock to avoid TOCTOU race
        now = _time.time()
        if now - _blacklist_cache["loaded_at"] < BLACKLIST_CACHE_TTL:
            return

        try:
            from auth_service import get_supabase_admin
            sb = get_supabase_admin()
            result = sb.table("blacklist").select("type, value").execute()
            emails, ips, domains = set(), set(), set()
            for row in (result.data or []):
                val = (row.get("value") or "").lower().strip()
                t = row.get("type", "")
                if t == "email":
                    emails.add(val)
                elif t == "ip":
                    ips.add(val)
                elif t == "domain":
                    domains.add(val)
            _blacklist_cache["emails"] = emails
            _blacklist_cache["ips"] = ips
            _blacklist_cache["domains"] = domains
            _blacklist_cache["loaded_at"] = now
        except Exception as e:
            logger.warning(f"Erro ao carregar blacklist (ignorado): {e}")


def _check_blacklist(request: Request, user: dict = None):
    """Verifica se o request vem de email/IP/domínio bloqueado."""
    _load_blacklist()

    # Verificar IP
    ip = get_remote_address(request)
    if ip in _blacklist_cache["ips"]:
        logger.warning(f"[BLACKLIST] IP bloqueado: {ip}")
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Acesso bloqueado. Contacte o administrador.",
        )

    # Verificar email e domínio (apenas se utilizador autenticado)
    email = ""
    if user:
        email = (user.get("email") or "").lower().strip()
    # Se não há utilizador autenticado, apenas verificação de IP (já feita acima)

    if email:
        if email in _blacklist_cache["emails"]:
            logger.warning(f"[BLACKLIST] Email bloqueado: {email[:3]}***")
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Conta bloqueada. Contacte o administrador.",
            )

        domain = email.split("@")[-1] if "@" in email else ""
        if domain in _blacklist_cache["domains"]:
            logger.warning(f"[BLACKLIST] Domínio bloqueado: {domain}")
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Domínio de email bloqueado. Contacte o administrador.",
            )


# ============================================================
# LIFESPAN - startup / shutdown
# ============================================================

def _cleanup_orphan_blocks():
    """Limpa créditos bloqueados há mais de 2 horas (órfãos de crashes/deploys)."""
    try:
        sb = get_supabase_admin()
        cutoff = (datetime.now(timezone.utc) - timedelta(hours=2)).isoformat()
        resp = sb.table("blocked_credits").select(
            "id, user_id, analysis_id, amount"
        ).eq("status", "blocked").lt("created_at", cutoff).execute()
        orphans = resp.data or []
        if not orphans:
            logger.info("[CLEANUP] Sem créditos órfãos bloqueados.")
            return
        logger.warning(f"[CLEANUP] {len(orphans)} bloqueio(s) órfão(s) encontrado(s) — a desbloquear...")
        from src.engine import cancelar_bloqueio
        for block in orphans:
            try:
                cancelar_bloqueio(block["analysis_id"])
                logger.info(f"[CLEANUP] Desbloqueado: analysis={block['analysis_id']}, ${block['amount']:.4f}")
            except Exception as e:
                logger.error(f"[CLEANUP] Falha ao desbloquear {block['analysis_id']}: {e}")
    except Exception as e:
        logger.warning(f"[CLEANUP] Erro ao verificar bloqueios órfãos: {e}")


def _sigterm_handler(signum, frame):
    """Handler SIGTERM/SIGINT: desbloquear créditos de análises activas antes de morrer."""
    sig_name = "SIGTERM" if signum == signal.SIGTERM else "SIGINT"
    logger.warning(f"[SHUTDOWN] {sig_name} recebido — a limpar créditos bloqueados...")
    with _active_lock:
        active_ids = list(_active_user_analyses.items())
    if not active_ids:
        logger.info("[SHUTDOWN] Sem análises activas — shutdown limpo.")
        sys.exit(0)
    logger.warning(f"[SHUTDOWN] {len(active_ids)} análise(s) activa(s) — a cancelar bloqueios...")
    try:
        from src.engine import cancelar_bloqueio
        for user_id, analysis_id in active_ids:
            if analysis_id and analysis_id != "starting":
                try:
                    cancelar_bloqueio(analysis_id)
                    logger.info(f"[SHUTDOWN] Bloqueio cancelado: user={user_id[:8]}, analysis={analysis_id}")
                except Exception as e:
                    logger.error(f"[SHUTDOWN] Falha ao cancelar bloqueio {analysis_id}: {e}")
            else:
                logger.warning(f"[SHUTDOWN] User {user_id[:8]} em fase 'starting' — sem analysis_id para cancelar")
    except Exception as e:
        logger.error(f"[SHUTDOWN] Erro geral no cleanup: {e}")
    sys.exit(0)


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Inicializa recursos no arranque e limpa no shutdown."""
    # -- Startup --
    supabase_url = os.environ.get("SUPABASE_URL", "")
    supabase_key = os.environ.get("SUPABASE_KEY", "")
    service_role_key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")

    if not supabase_url or not supabase_key:
        logger.warning("[AVISO] SUPABASE_URL ou SUPABASE_KEY não definidos no .env")
    else:
        get_supabase()
        logger.info(f"[OK] Supabase (anon) conectado: {supabase_url[:40]}...")

    if not service_role_key:
        logger.warning("[AVISO] SUPABASE_SERVICE_ROLE_KEY não definida - operações de servidor falharão")
    else:
        get_supabase_admin()
        logger.info("[OK] Supabase (service_role) conectado.")

    # v5.2: Registar signal handlers para cleanup de créditos
    signal.signal(signal.SIGTERM, _sigterm_handler)
    signal.signal(signal.SIGINT, _sigterm_handler)
    logger.info("[OK] Signal handlers (SIGTERM/SIGINT) registados.")

    # v5.2: Cleanup de créditos órfãos no startup
    _cleanup_orphan_blocks()

    logger.info("[OK] LexForum - Servidor iniciado.")
    yield
    # -- Shutdown --
    logger.info("[OK] Servidor encerrado.")


# ============================================================
# APP
# ============================================================

app = FastAPI(
    title="LexForum",
    description="Câmara de Análise Jurídica com IA",
    version="2.0.0",
    lifespan=lifespan,
)

# Rate limiter
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_handler)

# CORS - apenas origens autorizadas
CORS_ORIGINS = [
    "https://lexforum.eu",
    "https://www.lexforum.eu",
    "https://lovable.dev",
    "https://lexportal.lovable.app",
    "https://lexforum.lovable.app",
    "https://lexportal.lovable.dev",
]
if os.environ.get("ENV", "production").lower() in ("development", "dev", "local"):
    CORS_ORIGINS.extend([
        "http://localhost:3000",
        "http://localhost:5173",
        "http://localhost:8080",
    ])
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_origin_regex=r"https://[a-zA-Z0-9-]+\.lovable\.(app|dev)$",
    allow_credentials=True,
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type"],
)


# Security headers + blacklist check
@app.middleware("http")
async def security_middleware(request: Request, call_next):
    # Log Origin for CORS debugging
    origin = request.headers.get("origin", "NO-ORIGIN")
    method = request.method
    path = request.url.path
    if method == "OPTIONS":
        logger.info(f"[CORS-DEBUG] {method} {path} Origin={origin}")

    # Blacklist check (skip para /health e /docs)
    if path not in ("/health", "/docs", "/openapi.json", "/redoc"):
        try:
            _check_blacklist(request)
        except HTTPException as exc:
            return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})

    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    return response


# ============================================================
# ROTAS PÚBLICAS
# ============================================================

@app.get("/health")
async def health():
    """Rota de saúde - verifica se o servidor está online."""
    return {"status": "online"}


# ============================================================
# ROTAS PROTEGIDAS (requerem autenticação)
# ============================================================

@app.get("/me")
@limiter.limit("60/minute")
async def me(request: Request, user: dict = Depends(get_current_user)):
    """Retorna dados do utilizador autenticado (rota de teste)."""
    return {
        "user_id": user["id"],
        "email": user["email"],
    }


# Anti-duplicado: impede 2 análises do mesmo user ao mesmo tempo
_active_user_analyses: dict[str, str] = {}  # user_id -> analysis_id
_active_lock = threading.Lock()


@app.post("/analyze")
@limiter.limit("10/minute")
@limiter.limit("50/hour")
@limiter.limit("100/day")
async def analyze(
    request: Request,
    file: UploadFile = File(...),
    area_direito: str = Form("Civil"),  # Default to Civil law if not specified
    perguntas_raw: str = Form(""),
    titulo: str = Form(""),
    tier: str = Form("bronze"),
    user: dict = Depends(get_current_user),
):
    """
    Executa a análise jurídica completa de um documento.

    Parâmetros:
      - file: Ficheiro (PDF, DOCX, XLSX, TXT)
      - area_direito: Área do direito (Civil, Penal, Trabalho, etc.)
      - perguntas_raw: Perguntas separadas por ---
      - titulo: Título do projecto (opcional)
      - tier: Tier selecionado (bronze, silver, gold)

    Fluxo:
      1. Bloqueia créditos estimados (wallet)
      2. Executa pipeline de 4 fases
      3. Liquida créditos (custo real) ou cancela bloqueio (se erro)
      4. Retorna resultado completo em JSON
    """
    # Anti-duplo-clique: rejeitar se user já tem análise a correr
    user_id = user["id"]
    with _active_lock:
        if user_id in _active_user_analyses:
            existing = _active_user_analyses[user_id]
            logger.warning(f"[ANTI-DUP] User {user_id[:8]} já tem análise a correr: {existing}")
            raise HTTPException(
                status_code=status.HTTP_429_TOO_MANY_REQUESTS,
                detail="Já tem uma análise em curso. Aguarde que termine antes de submeter outra.",
            )
        _active_user_analyses[user_id] = "starting"

    try:
        # Pre-check Content-Length to reject oversized files early
        content_length = request.headers.get("content-length")
        if content_length:
            try:
                if int(content_length) > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                        detail="Ficheiro demasiado grande. Máximo: 50MB.",
                    )
            except ValueError:
                pass

        file_bytes = await file.read()

        if len(file_bytes) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail="Ficheiro demasiado grande. Máximo: 50MB.",
            )

        # Sanitize filename to prevent path traversal and prompt injection
        safe_filename = _sanitize_filename(file.filename)
        ext = os.path.splitext(safe_filename)[1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Tipo de ficheiro não suportado. Aceites: {', '.join(ALLOWED_EXTENSIONS)}",
            )

        logger.info(f"[TIER-DEBUG] Tier recebido do frontend: '{tier}'")

        # v4.0: Aceitar aliases standard/premium/elite
        tier_aliases = {
            "standard": "bronze",
            "premium": "silver",
            "elite": "gold",
        }
        tier = tier_aliases.get(tier.lower(), tier.lower())

        if tier not in ("bronze", "silver", "gold"):
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Tier inválido. Opções: bronze/standard, silver/premium, gold/elite.",
            )

        # v5.2: Gerar analysis_id no main para o SIGTERM handler poder cancelar bloqueios
        import uuid as _uuid
        _analysis_id = str(_uuid.uuid4())
        with _active_lock:
            _active_user_analyses[user_id] = _analysis_id

        # FIX 2026-02-14: Executar em thread separada para não bloquear event loop
        resultado = await asyncio.to_thread(
            executar_analise_documento,
            user_id=user["id"],
            file_bytes=file_bytes,
            filename=safe_filename,
            area_direito=area_direito,
            perguntas_raw=perguntas_raw,
            titulo=titulo,
            tier=tier,
            analysis_id=_analysis_id,
        )

        result_dict = resultado.to_dict()

        # Guardar resultado na tabela documents
        try:
            sb_admin = get_supabase_admin()
            custos = result_dict.get("custos") or {}
            custo_real = custos.get("custo_total_usd", 0)
            custo_cobrado = custos.get("custo_cliente_usd", 0)
            logger.debug(
                f"[DOCS-DEBUG] custos keys={list(custos.keys()) if custos else 'EMPTY'}, "
                f"custo_real={custo_real}, custo_cobrado={custo_cobrado}, "
                f"type_real={type(custo_real).__name__}, type_cobrado={type(custo_cobrado).__name__}"
            )
            doc_record = {
                "user_id": user["id"],
                "title": titulo or result_dict.get("documento_filename", ""),
                "analysis_result": result_dict,
                "status": "completed" if resultado.sucesso else "error",
                "tier": tier,
                "area_direito": area_direito.strip().title(),
                "run_id": result_dict.get("run_id", ""),
                "filename": safe_filename,
                "file_size_bytes": len(file_bytes),
                "total_tokens": result_dict.get("total_tokens", 0),
                "custo_real_usd": float(custo_real) if custo_real else 0.0,
                "custo_cobrado_usd": float(custo_cobrado) if custo_cobrado else 0.0,
                "duracao_segundos": result_dict.get("duracao_total_s", 0),
            }
            insert_resp = sb_admin.table("documents").insert(doc_record).execute()
            if not insert_resp.data:
                logger.warning("[DOCS] Insert returned empty data - document record may not have been saved")
            doc_id = insert_resp.data[0]["id"] if insert_resp.data else None
            if doc_id:
                result_dict["document_id"] = doc_id
                logger.info(f"[DOCS] Resultado guardado: doc_id={doc_id}")
                # Ligar document_id aos registos de model_performance
                try:
                    from src.performance_tracker import PerformanceTracker
                    tracker = PerformanceTracker.get_instance()
                    if tracker:
                        tracker.link_document_id(result_dict.get("run_id", ""), doc_id)
                except Exception as e:
                    logger.warning(f"[PERF] Falha ao ligar document_id: {e}")
        except Exception as e:
            logger.warning(f"[DOCS] Erro ao guardar resultado: {e}")

        return result_dict

    except InsufficientBalanceError as e:
        raise HTTPException(
            status_code=status.HTTP_402_PAYMENT_REQUIRED,
            detail={
                "message": "Saldo insuficiente. Por favor carregue a conta.",
                "saldo_atual": e.saldo_atual,
                "saldo_necessario": getattr(e, 'saldo_necessario', e.saldo_minimo),
                "moeda": "USD",
            },
        )
    except InvalidDocumentError as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=str(e),
        )
    except MissingApiKeyError as e:
        logger.error(f"API Key em falta: {e}")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Serviço temporariamente indisponível. Contacte o suporte.",
        )
    except BudgetExceededError as e:
        raise HTTPException(
            status_code=status.HTTP_402_PAYMENT_REQUIRED,
            detail={
                "error": "insufficient_credits",
                "message": str(e),
                "detail": "Saldo insuficiente. Recarregue a sua wallet para continuar.",
            },
        )
    except EngineError as e:
        logger.error(f"Erro do engine: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e),
        )
    except Exception:
        logger.exception("Erro inesperado no endpoint /analyze")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Erro interno do servidor.",
        )
    finally:
        # Anti-duplo-clique: libertar o user
        with _active_lock:
            _active_user_analyses.pop(user_id, None)


# ============================================================
# EXPORTAÇÃO DE RELATÓRIOS (PDF / DOCX)
# ============================================================

class ExportRequest(BaseModel):
    analysis_result: dict[str, Any]

    @field_validator("analysis_result")
    @classmethod
    def validate_size(cls, v):
        import json as _json
        if len(_json.dumps(v, default=str)) > 10 * 1024 * 1024:
            raise ValueError("analysis_result excede o tamanho máximo permitido (10MB)")
        return v

_INTERNAL_PATTERNS = [
    re.compile(r"^\s*-?\s*Fontes:\s*\[.*?\]\s*$", re.MULTILINE),
    re.compile(r"^\s*-?\s*Local:\s*chars\s+[\d,]+-[\d,]+\s*$", re.MULTILINE),
    re.compile(r"^\s*##?\s*RELATÓRIO DE COBERTURA.*$", re.MULTILINE),
    re.compile(r"^\s*##?\s*CONFLITOS DETETADOS.*$", re.MULTILINE),
    re.compile(r"^\s*##?\s*ERROS\s*$", re.MULTILINE),
    re.compile(r"^\s*-\s*\*\*Total chars:\*\*.*$", re.MULTILINE),
    re.compile(r"^\s*-\s*\*\*Chars cobertos:\*\*.*$", re.MULTILINE),
    re.compile(r"^\s*-\s*\*\*Completa:\*\*.*$", re.MULTILINE),
    re.compile(r"^.*INTEGRITY_WARNING.*$", re.MULTILINE),
    re.compile(r"^.*EXCERPT_MISMATCH.*$", re.MULTILINE),
    re.compile(r"^.*RANGE_INVALID.*$", re.MULTILINE),
    re.compile(r"^.*PAGE_MISMATCH.*$", re.MULTILINE),
    re.compile(r"^.*ITEM_NOT_FOUND.*$", re.MULTILINE),
    re.compile(r"^.*MISSING_CITATION.*$", re.MULTILINE),
    re.compile(r"^.*match_ratio=.*$", re.MULTILINE),
]


def _sanitize_content(text: str) -> str:
    """Remove metadados técnicos/internos do conteúdo antes de apresentar ao utilizador."""
    if not text:
        return text
    for pattern in _INTERNAL_PATTERNS:
        text = pattern.sub("", text)
    # Limpar linhas vazias consecutivas
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _build_pdf(data: dict[str, Any]) -> bytes:
    """Gera PDF profissional a partir do resultado da análise."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "LexForumTitle", parent=styles["Title"],
        fontSize=20, textColor=HexColor("#1a1a2e"), spaceAfter=20,
    ))
    styles.add(ParagraphStyle(
        "SectionHead", parent=styles["Heading1"],
        fontSize=14, textColor=HexColor("#16213e"),
        spaceBefore=16, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        "BodyText2", parent=styles["BodyText"],
        fontSize=10, leading=14, spaceAfter=6,
    ))

    elements = []

    # Título
    elements.append(Paragraph("Relatório de Análise — LexForum", styles["LexForumTitle"]))
    elements.append(Spacer(1, 10))

    # Metadados (sem custos — custos apenas no ecrã, nunca em exportação)
    meta = [
        ["Área de Direito", data.get("area_direito", "N/A")],
        ["Data da Análise", data.get("timestamp_inicio", "N/A")[:10] if data.get("timestamp_inicio") else "N/A"],
    ]
    t = Table(meta, colWidths=[5 * cm, 12 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), HexColor("#e8e8e8")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 16))

    # Conclusão (veredicto)
    simbolo = data.get("simbolo_final", "")
    veredicto = data.get("veredicto_final", "Sem conclusão")
    elements.append(Paragraph("Conclusão", styles["SectionHead"]))
    elements.append(Paragraph(f"{simbolo} {veredicto}", styles["BodyText2"]))
    elements.append(Spacer(1, 10))

    # Relatório Profissional (Curador Sénior — substitui fases técnicas)
    f4 = _sanitize_content(data.get("fase3_presidente", ""))
    for line in (f4 or "Sem análise disponível.").split("\n"):
        if line.strip():
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            elements.append(Paragraph(safe, styles["BodyText2"]))

    # Rodapé
    elements.append(Spacer(1, 30))
    elements.append(Paragraph(
        "Gerado automaticamente por LexForum — Este documento não substitui aconselhamento jurídico.",
        ParagraphStyle("Footer", parent=styles["Normal"], fontSize=7, textColor=HexColor("#999999")),
    ))

    doc.build(elements)
    return buf.getvalue()


def _build_docx(data: dict[str, Any]) -> bytes:
    """Gera DOCX profissional a partir do resultado da análise."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    title = doc.add_heading("Relatório de Análise — LexForum", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadados (sem custos — custos apenas no ecrã, nunca em exportação)
    meta_rows = [
        ("Área de Direito", data.get("area_direito", "N/A")),
        ("Data da Análise", data.get("timestamp_inicio", "N/A")[:10] if data.get("timestamp_inicio") else "N/A"),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(meta_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value) if value else "N/A"
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph("")

    doc.add_heading("Conclusão", level=1)
    simbolo = data.get("simbolo_final", "")
    veredicto = data.get("veredicto_final", "Sem conclusão")
    p = doc.add_paragraph(f"{simbolo} {veredicto}")
    if p.runs: p.runs[0].bold = True

    # Relatório Profissional (Curador Sénior — substitui fases técnicas)
    f4 = _sanitize_content(data.get("fase3_presidente", ""))
    doc.add_paragraph(f4 or "Sem análise disponível.")

    doc.add_paragraph("")
    footer = doc.add_paragraph(
        "Gerado automaticamente por LexForum — "
        "Este documento não substitui aconselhamento jurídico."
    )
    if footer.runs:
        footer.runs[0].font.size = Pt(7)
        footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post("/export/pdf")
@limiter.limit("30/minute")
async def export_pdf(request: Request, req: ExportRequest, user: dict = Depends(get_current_user)):
    """Exporta o resultado da análise como PDF."""
    try:
        pdf_bytes = _build_pdf(req.analysis_result)
        run_id = req.analysis_result.get("run_id", "relatorio")
        safe_id = "".join(c for c in str(run_id) if c.isalnum() or c in "-_")[:64]
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="relatorio_{safe_id}.pdf"'},
        )
    except Exception:
        logger.exception("Erro ao gerar PDF")
        raise HTTPException(status_code=500, detail="Erro ao gerar PDF.")


@app.post("/export/docx")
@limiter.limit("30/minute")
async def export_docx(request: Request, req: ExportRequest, user: dict = Depends(get_current_user)):
    """Exporta o resultado da análise como DOCX."""
    try:
        docx_bytes = _build_docx(req.analysis_result)
        run_id = req.analysis_result.get("run_id", "relatorio")
        safe_id = "".join(c for c in str(run_id) if c.isalnum() or c in "-_")[:64]
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="relatorio_{safe_id}.docx"'},
        )
    except Exception:
        logger.exception("Erro ao gerar DOCX")
        raise HTTPException(status_code=500, detail="Erro ao gerar DOCX.")


# ============================================================
# PERGUNTAS PÓS-ANÁLISE (POST /ask)
# ============================================================

from src.config import ASK_MODELS

ASK_SYSTEM_PROMPT = (
    "És um jurista especializado em Direito Português. "
    "Com base na análise jurídica fornecida, responde à pergunta do utilizador "
    "de forma clara e fundamentada, citando legislação quando aplicável."
)

CONSOLIDATION_SYSTEM_PROMPT = (
    "Recebeste 3 respostas de juristas diferentes à mesma pergunta. "
    "Consolida-as numa única resposta coerente, mantendo os pontos de consenso "
    "e assinalando divergências. Cita legislação mencionada. "
    "Responde em português de Portugal."
)


class AskRequest(BaseModel):
    question: str
    analysis_result: dict[str, Any]
    document_id: str = ""
    previous_qa: list[dict[str, str]] = Field(default=[], max_length=50)

    @field_validator("analysis_result")
    @classmethod
    def validate_size(cls, v):
        import json as _json
        if len(_json.dumps(v, default=str)) > 10 * 1024 * 1024:
            raise ValueError("analysis_result excede o tamanho máximo permitido (10MB)")
        return v


def _build_ask_context(data: dict[str, Any], previous_qa: list[dict[str, str]] = None) -> str:
    """Extrai contexto relevante do resultado da análise para a pergunta."""
    parts = []

    area = data.get("area_direito", "")
    if area:
        parts.append(f"Área de Direito: {area}")

    veredicto = data.get("veredicto_final", "")
    simbolo = data.get("simbolo_final", "")
    if veredicto:
        parts.append(f"Parecer: {simbolo} {veredicto}")

    f1 = data.get("fase1_agregado_consolidado") or data.get("fase1_agregado", "")
    if f1:
        parts.append(f"EXTRAÇÃO:\n{f1[:3000]}")

    f2 = data.get("fase2_chefe_consolidado") or data.get("fase2_chefe", "")
    if f2:
        parts.append(f"AUDITORIA:\n{f2[:3000]}")

    f4 = data.get("fase3_presidente", "")
    if f4:
        parts.append(f"ANÁLISE FINAL:\n{f4[:3000]}")

    docs_adicionais = data.get("documentos_adicionais", [])
    if docs_adicionais:
        parts.append("DOCUMENTOS ADICIONAIS:")
        for doc in docs_adicionais:
            nome = doc.get("filename", "documento")
            texto = doc.get("text", "")[:2000]
            parts.append(f"--- {nome} ---\n{texto}")

    if previous_qa:
        parts.append("PERGUNTAS E RESPOSTAS ANTERIORES:")
        for qa in previous_qa:
            parts.append(f"P: {qa.get('question', '')}")
            parts.append(f"R: {qa.get('answer', '')}")

    return "\n\n".join(parts) if parts else "Sem contexto de análise disponível."


@app.post("/ask")
@limiter.limit("20/minute")
@limiter.limit("100/hour")
async def ask_question(request: Request, req: AskRequest, user: dict = Depends(get_current_user)):
    """
    Pergunta pós-análise: envia a pergunta a 3 LLMs com contexto da análise
    anterior + histórico de Q&A e consolida as respostas.
    """
    from src.llm_client import get_llm_client

    question = req.question.strip()
    if not question:
        raise HTTPException(status_code=422, detail="A pergunta não pode estar vazia.")
    if len(question) > MAX_QUESTION_LENGTH:
        raise HTTPException(status_code=422, detail=f"Pergunta demasiado longa. Máximo: {MAX_QUESTION_LENGTH} caracteres.")

    context = _build_ask_context(req.analysis_result, req.previous_qa)

    prompt = f"""ANÁLISE JURÍDICA ANTERIOR:
{context}

PERGUNTA DO UTILIZADOR:
{question}

Responde de forma clara, citando legislação quando aplicável."""

    llm = get_llm_client()
    individual_responses: list[dict[str, str]] = []

    for model in ASK_MODELS:
        try:
            resp = await asyncio.to_thread(
                llm.chat_simple,
                model=model,
                prompt=prompt,
                system_prompt=ASK_SYSTEM_PROMPT,
                temperature=0.3,
                max_tokens=4096,
            )
            individual_responses.append({
                "model": model,
                "response": resp.content,
            })
        except Exception as e:
            logger.warning(f"Modelo {model} falhou no /ask: {e}")
            individual_responses.append({
                "model": model,
                "response": "[Erro: modelo indisponível]",
            })

    if len([r for r in individual_responses if not r["response"].startswith("[Erro")]) == 0:
        raise HTTPException(
            status_code=503,
            detail="Nenhum modelo conseguiu responder. Tente novamente.",
        )

    respostas_texto = "\n\n".join([
        f"## {r['model']}:\n{r['response']}"
        for r in individual_responses
        if not r["response"].startswith("[Erro")
    ])

    consolidation_prompt = f"""PERGUNTA: {question}

RESPOSTAS DOS 3 JURISTAS:
{respostas_texto}

Consolida estas respostas numa única resposta final coerente."""

    try:
        consolidated = await asyncio.to_thread(
            llm.chat_simple,
            model=ASK_MODELS[0],
            prompt=consolidation_prompt,
            system_prompt=CONSOLIDATION_SYSTEM_PROMPT,
            temperature=0.2,
            max_tokens=4096,
        )
        answer = consolidated.content
    except Exception as e:
        logger.warning(f"Consolidação falhou: {e}")
        answer = next(
            (r["response"] for r in individual_responses if not r["response"].startswith("[Erro")),
            "Não foi possível gerar resposta.",
        )

    if req.document_id:
        try:
            sb = get_supabase_admin()
            doc_resp = sb.table("documents").select("analysis_result").eq("id", req.document_id).eq("user_id", user["id"]).single().execute()
            current_result = (doc_resp.data or {}).get("analysis_result") or {}
            qa_history = current_result.get("qa_history", [])
            qa_history.append({
                "question": question,
                "answer": answer,
                "individual_responses": individual_responses,
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })
            current_result["qa_history"] = qa_history
            sb.table("documents").update(
                {"analysis_result": current_result}
            ).eq("id", req.document_id).eq("user_id", user["id"]).execute()
            logger.info(f"Q&A guardada no Supabase: doc={req.document_id}, total_qa={len(qa_history)}")
        except Exception as e:
            logger.warning(f"Erro ao guardar Q&A no Supabase: {e}")
            return {
                "question": question,
                "answer": answer,
                "individual_responses": individual_responses,
                "persisted": False,
            }

    return {
        "question": question,
        "answer": answer,
        "individual_responses": individual_responses,
    }


# ============================================================
# ADICIONAR DOCUMENTO A PROJECTO EXISTENTE
# ============================================================

@app.post("/analyze/add")
@limiter.limit("10/minute")
async def add_document_to_project(
    request: Request,
    file: UploadFile = File(...),
    document_id: str = Form(...),
    user: dict = Depends(get_current_user),
):
    """
    Adiciona um novo documento a um projecto existente.
    Extrai o texto e guarda-o no analysis_result do documento original.
    """
    from src.engine import carregar_documento_de_bytes

    file_bytes = await file.read()
    filename = _sanitize_filename(file.filename)

    # Validar extensão
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Tipo de ficheiro não suportado. Aceites: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="Ficheiro demasiado grande. Máximo: 50MB.",
        )

    if len(file_bytes) == 0:
        raise HTTPException(status_code=422, detail="Ficheiro vazio.")

    try:
        doc = carregar_documento_de_bytes(file_bytes, filename)
        novo_texto = doc.text
    except Exception:
        raise HTTPException(status_code=422, detail="Erro ao extrair texto do documento.")

    if not novo_texto or len(novo_texto.strip()) < 50:
        raise HTTPException(
            status_code=422,
            detail="O documento não contém texto suficiente.",
        )

    try:
        sb = get_supabase_admin()
        doc_resp = sb.table("documents").select("analysis_result, user_id").eq("id", document_id).eq("user_id", user["id"]).single().execute()
    except Exception:
        raise HTTPException(status_code=404, detail="Documento não encontrado.")

    if not doc_resp.data:
        raise HTTPException(status_code=404, detail="Documento não encontrado.")

    doc_user_id = doc_resp.data.get("user_id", "")
    if doc_user_id != user.get("id"):
        raise HTTPException(status_code=403, detail="Sem permissão para este documento.")

    current_result = doc_resp.data.get("analysis_result") or {}
    docs_adicionais = current_result.get("documentos_adicionais", [])

    if len(docs_adicionais) >= MAX_ADDITIONAL_DOCS:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Limite de {MAX_ADDITIONAL_DOCS} documentos adicionais atingido.",
        )

    docs_adicionais.append({
        "filename": filename,
        "text": novo_texto,
        "chars": len(novo_texto),
        "added_at": datetime.now(timezone.utc).isoformat(),
    })
    current_result["documentos_adicionais"] = docs_adicionais

    try:
        sb.table("documents").update(
            {"analysis_result": current_result}
        ).eq("id", document_id).execute()
    except Exception:
        raise HTTPException(status_code=500, detail="Erro ao guardar documento.")

    logger.info(
        f"Documento adicionado: doc={document_id}, file={filename}, "
        f"chars={len(novo_texto)}, total_docs_adicionais={len(docs_adicionais)}"
    )

    return {
        "status": "ok",
        "document_id": document_id,
        "added_file": filename,
        "added_chars": len(novo_texto),
        "total_additional_docs": len(docs_adicionais),
        "message": f"Documento '{filename}' adicionado ao projecto. "
                   f"Use /ask para fazer perguntas sobre todos os documentos.",
    }


# ============================================================
# TIER CONFIGURATION ENDPOINTS
# ============================================================

@app.get("/tiers")
@limiter.limit("60/minute")
async def get_tiers(request: Request):
    """Retorna configuração de todos os tiers (Bronze/Prata/Ouro)."""
    from src.tier_config import get_all_tiers_info
    return {"tiers": get_all_tiers_info()}


@app.post("/tiers/calculate")
@limiter.limit("60/minute")
async def calculate_tier_cost_endpoint(
    request: Request,
    tier: str,
    document_tokens: int = Query(default=0, ge=0, le=10_000_000),
):
    """
    Calcula custo estimado para um tier.
    Query params:
      tier: bronze, silver ou gold
      document_tokens: tamanho do documento em tokens (opcional)
    """
    from src.tier_config import TierLevel, calculate_tier_cost as calc_cost
    # v4.0: Aceitar aliases
    tier_map = {"standard": "bronze", "premium": "silver", "elite": "gold"}
    tier_normalized = tier_map.get(tier.lower(), tier.lower())
    try:
        tier_level = TierLevel(tier_normalized)
    except ValueError:
        raise HTTPException(
            status_code=422,
            detail=f"Tier inválido: {tier}. Use: bronze/standard, silver/premium, gold/elite"
        )

    costs = calc_cost(tier_level, document_tokens)
    return {
        "tier": tier,
        "custo_real_usd": costs["custo_real"],
        "custo_cliente_usd": costs["custo_cliente"],
        "bloqueio_usd": costs["bloqueio"],
        "size_multiplier": costs["size_multiplier"],
    }


# ============================================================
# WALLET ENDPOINTS
# ============================================================

@app.get("/wallet/balance")
@limiter.limit("60/minute")
async def wallet_balance(request: Request, user: dict = Depends(get_current_user)):
    """Retorna saldo actual da wallet do utilizador."""
    try:
        wm = get_wallet_manager()
        saldo = wm.get_balance(user["id"], user_email=user.get("email", ""))
        markup = wm.get_markup_multiplier()
        total = float(saldo.get("total", 0))
        blocked = float(saldo.get("blocked", 0))
        available = float(saldo.get("available", 0))
        return {
            "balance_usd": available,
            "total_usd": total,
            "blocked_usd": blocked,
            "moeda": "USD",
            "markup_multiplier": markup,
        }
    except Exception as e:
        logger.error(f"Erro ao consultar saldo: {e}")
        raise HTTPException(status_code=500, detail="Erro ao consultar saldo.")


@app.get("/wallet/transactions")
@limiter.limit("60/minute")
async def wallet_transactions(
    request: Request,
    limit: int = 50,
    offset: int = 0,
    type_filter: Optional[str] = None,
    user: dict = Depends(get_current_user),
):
    """Retorna histórico de transações da wallet."""
    try:
        wm = get_wallet_manager()
        return wm.get_transactions(
            user_id=user["id"],
            limit=min(limit, 100),
            offset=offset,
            type_filter=type_filter,
        )
    except Exception as e:
        logger.error(f"Erro ao consultar transações: {e}")
        raise HTTPException(status_code=500, detail="Erro ao consultar transações.")


class CreditRequest(BaseModel):
    user_id: str
    amount_usd: float
    description: str = ""


@app.post("/wallet/credit")
@limiter.limit("10/minute")
@limiter.limit("30/hour")
async def wallet_credit(
    request: Request,
    req: CreditRequest,
    user: dict = Depends(get_current_user),
):
    """Credita saldo na wallet de um utilizador (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Apenas administradores podem creditar saldos.",
        )

    if req.amount_usd <= 0:
        raise HTTPException(status_code=422, detail="Valor deve ser positivo.")

    try:
        wm = get_wallet_manager()
        result = wm.credit(
            user_id=req.user_id,
            amount_usd=req.amount_usd,
            description=req.description,
            admin_id=user["id"],
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        logger.error(f"Erro ao creditar saldo: {e}")
        raise HTTPException(status_code=500, detail="Erro ao creditar saldo.")


# ============================================================
# ADMIN - LISTA DE UTILIZADORES
# ============================================================

@app.get("/admin/users")
@limiter.limit("30/minute")
async def admin_list_users(
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Lista todos os utilizadores com saldos (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(status_code=403, detail="Apenas administradores.")

    try:
        sb = get_supabase_admin()

        # Buscar todos os profiles com saldo
        profiles_resp = sb.table("profiles").select(
            "id, email, full_name, credits_balance, credits_blocked, created_at"
        ).order("credits_balance", desc=True).limit(500).execute()

        users = []
        for p in (profiles_resp.data or []):
            email = (p.get("email") or "").strip()
            # Ignorar perfis sem email (utilizadores de teste antigos)
            if not email:
                continue
            total = float(p.get("credits_balance") or 0)
            blocked = float(p.get("credits_blocked") or 0)
            users.append({
                "user_id": p["id"],
                "email": email,
                "full_name": p.get("full_name", ""),
                "balance_usd": total,
                "blocked_usd": blocked,
                "available_usd": total - blocked,
                "created_at": p.get("created_at", ""),
            })

        return {"users": users, "total": len(users)}

    except Exception as e:
        logger.error(f"Erro ao listar utilizadores: {e}")
        raise HTTPException(status_code=500, detail="Erro ao listar utilizadores.")


# ============================================================
# ADMIN - MODEL PERFORMANCE
# ============================================================

@app.get("/admin/model-performance")
@limiter.limit("30/minute")
async def admin_model_performance(
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Dashboard de performance por modelo de IA (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(status_code=403, detail="Apenas administradores.")

    try:
        sb = get_supabase_admin()

        # Stats agregados por modelo+role
        summary_resp = sb.table("model_performance").select(
            "model, role, success, excerpt_mismatches, range_invalids, "
            "page_mismatches, missing_citations, error_recovered, "
            "latency_ms, total_tokens, cost_usd, error_type"
        ).order("created_at", desc=True).limit(2000).execute()

        rows = summary_resp.data or []

        # Agregar por modelo
        from collections import defaultdict
        agg = defaultdict(lambda: {
            "calls": 0, "success": 0, "failed": 0,
            "total_tokens": 0, "total_cost": 0.0, "total_latency": 0.0,
            "excerpt_mismatches": 0, "range_invalids": 0,
            "page_mismatches": 0, "missing_citations": 0,
            "error_types": defaultdict(int),
        })
        for r in rows:
            key = f"{r['model']}|{r['role']}"
            a = agg[key]
            a["model"] = r["model"]
            a["role"] = r["role"]
            a["calls"] += 1
            if r["success"]:
                a["success"] += 1
            else:
                a["failed"] += 1
                if r.get("error_type"):
                    a["error_types"][r["error_type"]] += 1
            a["total_tokens"] += r.get("total_tokens") or 0
            a["total_cost"] += float(r.get("cost_usd") or 0)
            a["total_latency"] += float(r.get("latency_ms") or 0)
            a["excerpt_mismatches"] += r.get("excerpt_mismatches") or 0
            a["range_invalids"] += r.get("range_invalids") or 0
            a["page_mismatches"] += r.get("page_mismatches") or 0
            a["missing_citations"] += r.get("missing_citations") or 0

        models = []
        for _key, a in sorted(agg.items(), key=lambda x: -x[1]["failed"]):
            tc = max(a["calls"], 1)
            models.append({
                "model": a.get("model", "?"),
                "role": a.get("role", "?"),
                "total_calls": a["calls"],
                "successful": a["success"],
                "failed": a["failed"],
                "success_rate": round(a["success"] / tc * 100, 1),
                "avg_tokens": round(a["total_tokens"] / tc),
                "avg_latency_ms": round(a["total_latency"] / tc, 1),
                "avg_cost_usd": round(a["total_cost"] / tc, 6),
                "total_cost_usd": round(a["total_cost"], 4),
                "excerpt_mismatches": a["excerpt_mismatches"],
                "range_invalids": a["range_invalids"],
                "page_mismatches": a["page_mismatches"],
                "missing_citations": a["missing_citations"],
                "error_types": dict(a["error_types"]),
            })

        # Piores performers (por taxa de erro)
        worst = [m for m in models if m["failed"] > 0]
        worst.sort(key=lambda x: x["success_rate"])

        # Ultimos erros
        errors_resp = sb.table("model_performance").select(
            "model, role, error_type, error_message, created_at"
        ).eq("success", False).order(
            "created_at", desc=True
        ).limit(50).execute()

        return {
            "total_records": len(rows),
            "models": models,
            "worst_performers": worst[:10],
            "recent_errors": errors_resp.data or [],
        }
    except Exception as e:
        logger.error(f"Erro em model-performance: {e}")
        raise HTTPException(status_code=500, detail="Erro ao consultar performance.")


# ============================================================
# ADMIN - DIAGNOSTICO TECNICO COMPLETO
# ============================================================

_DIAG_PHASE_LABELS = {
    "phase0": "Fase 0 \u2013 Triagem",
    "phase1": "Fase 1 \u2013 Extrac\u00e7\u00e3o",
    "phase2": "Fase 2 \u2013 Agrega\u00e7\u00e3o",
    "phase3": "Fase 3 \u2013 Auditoria",
    "phase4": "Fase 4 \u2013 Julgamento",
    "phase5": "Fase 5 \u2013 S\u00edntese",
}


def _diag_phase_number(phase_str: str) -> str:
    """Extrai 'phase1' de 'fase1_E3'."""
    if not phase_str:
        return "unknown"
    part = phase_str.split("_")[0]
    num = "".join(c for c in part if c.isdigit())
    return f"phase{num}" if num else "unknown"


def _diag_system_health() -> dict:
    """Recolhe estado do sistema (circuit breaker, pricing, analises activas)."""
    circuit = {"openai_open": False, "reason": "", "opened_at": None}
    try:
        from src.llm_client import get_llm_client
        client = get_llm_client()
        with client._circuit_lock:
            circuit["openai_open"] = client._openai_circuit_open
            circuit["reason"] = client._openai_circuit_reason or ""
            if client._openai_circuit_opened_at:
                circuit["opened_at"] = client._openai_circuit_opened_at.isoformat()
    except Exception:
        pass

    pricing_source = "unknown"
    pricing_cache = {}
    try:
        from src.cost_controller import DynamicPricing
        pricing_source = DynamicPricing.get_pricing_source()
        pricing_cache = DynamicPricing.get_cache_info()
    except Exception:
        pass

    with _active_lock:
        active_count = len(_active_user_analyses)

    return {
        "circuit_breaker": circuit,
        "pricing_source": pricing_source,
        "pricing_cache": pricing_cache,
        "active_analyses": active_count,
    }


def _diag_per_phase(rows: list) -> list:
    """Agrega metricas por fase do pipeline."""
    from collections import defaultdict
    phases: dict = defaultdict(lambda: {
        "calls": 0, "success": 0, "total_latency": 0.0, "total_cost": 0.0,
        "prompt_tokens": 0, "completion_tokens": 0, "cached_tokens": 0, "reasoning_tokens": 0,
    })
    for r in rows:
        p = _diag_phase_number(r.get("phase", ""))
        ph = phases[p]
        ph["calls"] += 1
        if r.get("success"):
            ph["success"] += 1
        ph["total_latency"] += float(r.get("latency_ms") or 0)
        ph["total_cost"] += float(r.get("cost_usd") or 0)
        ph["prompt_tokens"] += int(r.get("prompt_tokens") or 0)
        ph["completion_tokens"] += int(r.get("completion_tokens") or 0)
        ph["cached_tokens"] += int(r.get("cached_tokens") or 0)
        ph["reasoning_tokens"] += int(r.get("reasoning_tokens") or 0)

    total_latency_all = sum(ph["total_latency"] for ph in phases.values()) or 1.0
    max_avg_lat = 0.0
    max_lat_phase = ""
    result = []
    for p_key in sorted(phases.keys()):
        ph = phases[p_key]
        tc = max(ph["calls"], 1)
        avg_lat = ph["total_latency"] / tc
        if ph["calls"] >= 3 and avg_lat > max_avg_lat:
            max_avg_lat = avg_lat
            max_lat_phase = p_key
        result.append({
            "phase": p_key,
            "phase_label": _DIAG_PHASE_LABELS.get(p_key, p_key),
            "total_calls": ph["calls"],
            "success_rate": round(ph["success"] / tc * 100, 1),
            "avg_latency_ms": round(avg_lat),
            "avg_cost_usd": round(ph["total_cost"] / tc, 4),
            "total_cost_usd": round(ph["total_cost"], 4),
            "total_prompt_tokens": ph["prompt_tokens"],
            "total_completion_tokens": ph["completion_tokens"],
            "total_cached_tokens": ph["cached_tokens"],
            "total_reasoning_tokens": ph["reasoning_tokens"],
            "is_bottleneck": False,
            "pct_of_total_time": round(ph["total_latency"] / total_latency_all * 100, 1),
        })
    for item in result:
        if item["phase"] == max_lat_phase:
            item["is_bottleneck"] = True
    return result


def _diag_per_model(rows: list) -> list:
    """Agrega metricas por modelo de IA."""
    from collections import defaultdict
    models: dict = defaultdict(lambda: {
        "roles": set(), "calls": 0, "success": 0,
        "total_latency": 0.0, "total_cost": 0.0, "total_tokens": 0,
        "cached_tokens": 0, "prompt_tokens": 0,
        "excerpt_mismatches": 0, "range_invalids": 0,
        "page_mismatches": 0, "missing_citations": 0,
        "confidence_sum": 0.0, "confidence_count": 0,
        "errors": defaultdict(int),
        "retries": 0, "retry_success": 0,
    })
    for r in rows:
        m = r.get("model") or "unknown"
        md = models[m]
        md["roles"].add(r.get("role") or "?")
        md["calls"] += 1
        if r.get("success"):
            md["success"] += 1
        md["total_latency"] += float(r.get("latency_ms") or 0)
        md["total_cost"] += float(r.get("cost_usd") or 0)
        md["total_tokens"] += int(r.get("total_tokens") or 0)
        md["cached_tokens"] += int(r.get("cached_tokens") or 0)
        md["prompt_tokens"] += int(r.get("prompt_tokens") or 0)
        md["excerpt_mismatches"] += int(r.get("excerpt_mismatches") or 0)
        md["range_invalids"] += int(r.get("range_invalids") or 0)
        md["page_mismatches"] += int(r.get("page_mismatches") or 0)
        md["missing_citations"] += int(r.get("missing_citations") or 0)
        conf = r.get("confidence_adjusted")
        if conf is not None:
            md["confidence_sum"] += float(conf)
            md["confidence_count"] += 1
        if not r.get("success") and r.get("error_type"):
            md["errors"][r["error_type"]] += 1
        if r.get("was_retry"):
            md["retries"] += 1
            if r.get("success"):
                md["retry_success"] += 1

    result = []
    for m_name in sorted(models.keys()):
        md = models[m_name]
        tc = max(md["calls"], 1)
        tt = max(md["total_tokens"], 1)
        result.append({
            "model": m_name,
            "roles": sorted(md["roles"]),
            "total_calls": md["calls"],
            "successful": md["success"],
            "failed": md["calls"] - md["success"],
            "success_rate": round(md["success"] / tc * 100, 1),
            "avg_latency_ms": round(md["total_latency"] / tc),
            "avg_cost_usd": round(md["total_cost"] / tc, 4),
            "total_cost_usd": round(md["total_cost"], 4),
            "cost_per_1k_tokens": round(md["total_cost"] / tt * 1000, 4),
            "cache_hit_rate": round(md["cached_tokens"] / max(md["prompt_tokens"], 1) * 100, 1),
            "quality": {
                "excerpt_mismatches": md["excerpt_mismatches"],
                "range_invalids": md["range_invalids"],
                "page_mismatches": md["page_mismatches"],
                "missing_citations": md["missing_citations"],
                "avg_confidence": round(
                    md["confidence_sum"] / md["confidence_count"], 2
                ) if md["confidence_count"] > 0 else None,
            },
            "errors": dict(md["errors"]),
            "retry_stats": {
                "total_retries": md["retries"],
                "retry_success": md["retry_success"],
                "retry_success_rate": round(
                    md["retry_success"] / max(md["retries"], 1) * 100, 1
                ),
            },
        })
    result.sort(key=lambda x: x["total_calls"], reverse=True)
    return result


def _diag_quality(rows: list) -> dict:
    """Agrega metricas de qualidade (citacoes, confianca, hints)."""
    from collections import Counter
    total_findings = 0
    total_citations = 0
    excerpt_total = 0
    excerpt_mismatches = 0
    page_total = 0
    page_mismatches = 0
    range_total = 0
    range_invalids = 0
    conf_buckets = {"high_90_100": 0, "good_75_90": 0, "moderate_50_75": 0, "low_0_50": 0}
    hints_counter: Counter = Counter()

    for r in rows:
        fc = int(r.get("findings_count") or 0)
        cc = int(r.get("citations_count") or 0)
        total_findings += fc
        total_citations += cc
        em = int(r.get("excerpt_mismatches") or 0)
        if cc > 0:
            excerpt_total += cc
            excerpt_mismatches += em
        pm = int(r.get("page_mismatches") or 0)
        if cc > 0:
            page_total += cc
            page_mismatches += pm
        rv = int(r.get("range_invalids") or 0)
        if cc > 0:
            range_total += cc
            range_invalids += rv
        conf = r.get("confidence_adjusted")
        if conf is not None:
            c = float(conf)
            if c >= 0.9:
                conf_buckets["high_90_100"] += 1
            elif c >= 0.75:
                conf_buckets["good_75_90"] += 1
            elif c >= 0.5:
                conf_buckets["moderate_50_75"] += 1
            else:
                conf_buckets["low_0_50"] += 1
        hints = r.get("adaptive_hints_used")
        if hints and isinstance(hints, list):
            for h in hints:
                hints_counter[h] += 1

    return {
        "citation_accuracy": {
            "total_findings": total_findings,
            "total_citations": total_citations,
            "avg_citations_per_finding": round(
                total_citations / max(total_findings, 1), 2
            ),
            "findings_without_citations_pct": round(
                max(0, total_findings - total_citations) / max(total_findings, 1) * 100, 1
            ) if total_findings > total_citations else 0.0,
        },
        "excerpt_match_rate": round(
            (1 - excerpt_mismatches / max(excerpt_total, 1)) * 100, 1
        ),
        "page_match_rate": round(
            (1 - page_mismatches / max(page_total, 1)) * 100, 1
        ),
        "range_valid_rate": round(
            (1 - range_invalids / max(range_total, 1)) * 100, 1
        ),
        "confidence_distribution": conf_buckets,
        "top_adaptive_hints": [
            {"hint": h, "count": c} for h, c in hints_counter.most_common(10)
        ],
    }


def _diag_costs(rows: list, transactions: list) -> dict:
    """Agrega custos por fase e tendencia diaria."""
    from collections import defaultdict
    cost_by_phase: dict = defaultdict(float)
    for r in rows:
        p = _diag_phase_number(r.get("phase", ""))
        cost_by_phase[p] += float(r.get("cost_usd") or 0)

    daily: dict = defaultdict(lambda: {"cost_usd": 0.0, "charged_usd": 0.0, "analyses": 0})
    for tx in transactions:
        d = (tx.get("created_at") or "")[:10]
        if d:
            daily[d]["cost_usd"] += float(tx.get("cost_real_usd") or 0)
            daily[d]["charged_usd"] += float(tx.get("amount_usd") or 0)
            daily[d]["analyses"] += 1

    total_cost = sum(cost_by_phase.values())
    total_charged = sum(dd["charged_usd"] for dd in daily.values())
    total_analyses = sum(dd["analyses"] for dd in daily.values()) or 1

    return {
        "cost_per_phase": sorted(
            [
                {
                    "phase": p,
                    "label": _DIAG_PHASE_LABELS.get(p, p),
                    "cost_usd": round(c, 4),
                }
                for p, c in cost_by_phase.items()
            ],
            key=lambda x: x["cost_usd"],
            reverse=True,
        ),
        "daily_trend": sorted(
            [
                {
                    "date": d,
                    "cost_usd": round(v["cost_usd"], 4),
                    "charged_usd": round(v["charged_usd"], 4),
                    "analyses": v["analyses"],
                }
                for d, v in daily.items()
            ],
            key=lambda x: x["date"],
        ),
        "total_cost_usd": round(total_cost, 4),
        "total_charged_usd": round(total_charged, 4),
        "markup_effective_pct": round(
            (total_charged / max(total_cost, 0.0001) - 1) * 100, 1
        ) if total_cost > 0 else 0.0,
        "avg_cost_per_analysis": round(total_cost / total_analyses, 4),
        "avg_charge_per_analysis": round(total_charged / total_analyses, 4),
    }


def _diag_errors(rows: list) -> dict:
    """Agrega erros por tipo, tendencia e piores modelos."""
    from collections import defaultdict, Counter
    error_types: Counter = Counter()
    daily_errors: dict = defaultdict(lambda: {"errors": 0, "total": 0})
    model_fails: dict = defaultdict(lambda: {"calls": 0, "fails": 0})
    total_retries = 0
    retry_success = 0
    total_errors = 0

    for r in rows:
        d = (r.get("created_at") or "")[:10]
        if d:
            daily_errors[d]["total"] += 1
        m = r.get("model") or "unknown"
        model_fails[m]["calls"] += 1
        if not r.get("success"):
            total_errors += 1
            et = r.get("error_type") or "OTHER"
            error_types[et] += 1
            if d:
                daily_errors[d]["errors"] += 1
            model_fails[m]["fails"] += 1
        if r.get("was_retry"):
            total_retries += 1
            if r.get("success"):
                retry_success += 1

    total_calls = max(len(rows), 1)

    worst = sorted(
        [
            {
                "model": m,
                "success_rate": round(
                    (1 - v["fails"] / max(v["calls"], 1)) * 100, 1
                ),
                "total_calls": v["calls"],
                "fails": v["fails"],
            }
            for m, v in model_fails.items()
            if v["calls"] >= 3
        ],
        key=lambda x: x["success_rate"],
    )

    return {
        "total_errors": total_errors,
        "error_rate_pct": round(total_errors / total_calls * 100, 1),
        "errors_by_type": [
            {"type": t, "count": c, "pct": round(c / max(total_errors, 1) * 100, 1)}
            for t, c in error_types.most_common()
        ],
        "daily_error_trend": sorted(
            [
                {
                    "date": d,
                    "errors": v["errors"],
                    "total_calls": v["total"],
                    "error_rate": round(v["errors"] / max(v["total"], 1) * 100, 1),
                }
                for d, v in daily_errors.items()
            ],
            key=lambda x: x["date"],
        ),
        "retry_stats": {
            "total_retries": total_retries,
            "successful_retries": retry_success,
            "retry_success_rate": round(
                retry_success / max(total_retries, 1) * 100, 1
            ),
        },
        "worst_models": worst[:5],
    }


def _diag_recent(docs: list) -> list:
    """Formata analises recentes."""
    result = []
    for d in docs:
        result.append({
            "id": d.get("id"),
            "title": d.get("title") or "Sem titulo",
            "status": d.get("status") or "unknown",
            "tier": d.get("tier") or "",
            "area": d.get("area_direito") or "",
            "cost_real_usd": float(d.get("custo_real_usd") or 0),
            "cost_charged_usd": float(d.get("custo_cobrado_usd") or 0),
            "duration_seconds": float(d.get("duracao_segundos") or 0),
            "tokens": int(d.get("total_tokens") or 0),
            "created_at": d.get("created_at") or "",
        })
    return result


@app.get("/admin/diagnostics")
@limiter.limit("10/minute")
async def admin_diagnostics(
    request: Request,
    days: int = Query(default=30, ge=1, le=90),
    user: dict = Depends(get_current_user),
):
    """Painel de diagnostico tecnico completo (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(status_code=403, detail="Apenas administradores.")

    try:
        sb = get_supabase_admin()
        cutoff = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()

        perf_resp = sb.table("model_performance").select(
            "phase, model, role, success, latency_ms, cost_usd, "
            "prompt_tokens, completion_tokens, total_tokens, cached_tokens, reasoning_tokens, "
            "excerpt_mismatches, range_invalids, page_mismatches, missing_citations, "
            "error_type, was_retry, "
            "confidence_adjusted, "
            "adaptive_hints_used, "
            "findings_count, citations_count, created_at"
        ).gte("created_at", cutoff).order("created_at", desc=True).limit(5000).execute()

        docs_resp = sb.table("documents").select(
            "id, title, status, tier, area_direito, custo_real_usd, "
            "custo_cobrado_usd, duracao_segundos, total_tokens, created_at"
        ).order("created_at", desc=True).limit(20).execute()

        tx_resp = sb.table("wallet_transactions").select(
            "amount_usd, cost_real_usd, created_at"
        ).eq("type", "debit").gte("created_at", cutoff).execute()

        rows = perf_resp.data or []

        return {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "data_window_days": days,
            "total_records": len(rows),
            "system_health": _diag_system_health(),
            "per_phase": _diag_per_phase(rows),
            "per_model": _diag_per_model(rows),
            "quality_metrics": _diag_quality(rows),
            "cost_analysis": _diag_costs(rows, tx_resp.data or []),
            "error_analysis": _diag_errors(rows),
            "recent_analyses": _diag_recent(docs_resp.data or []),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erro em diagnostics: {e}")
        raise HTTPException(status_code=500, detail="Erro ao gerar diagnostico.")


@app.get("/admin/profit-report")
@limiter.limit("30/minute")
async def admin_profit_report(
    request: Request,
    days: int = 30,
    user: dict = Depends(get_current_user),
):
    """Relatório de lucro (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Apenas administradores podem ver relatórios de lucro.",
        )

    try:
        wm = get_wallet_manager()
        return wm.get_profit_report(days=min(days, 365))
    except Exception as e:
        logger.error(f"Erro ao gerar relatório de lucro: {e}")
        raise HTTPException(status_code=500, detail="Erro ao gerar relatório.")


# ============================================================
# ADMIN - BLACKLIST
# ============================================================

class BlacklistAddRequest(BaseModel):
    type: str       # "email", "ip", ou "domain"
    value: str      # o email, IP ou domínio a bloquear
    reason: str = ""


@app.get("/admin/blacklist")
@limiter.limit("30/minute")
async def admin_blacklist_list(
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Lista todas as entradas na blacklist (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(status_code=403, detail="Apenas administradores.")

    try:
        sb = get_supabase_admin()
        result = sb.table("blacklist").select("*").order("created_at", desc=True).execute()
        return {"blacklist": result.data or [], "total": len(result.data or [])}
    except Exception as e:
        logger.error(f"Erro ao consultar blacklist: {e}")
        raise HTTPException(status_code=500, detail="Erro ao consultar blacklist.")


@app.post("/admin/blacklist")
@limiter.limit("30/minute")
async def admin_blacklist_add(
    req: BlacklistAddRequest,
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Adiciona email, IP ou domínio à blacklist (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(status_code=403, detail="Apenas administradores.")

    if req.type not in ("email", "ip", "domain"):
        raise HTTPException(status_code=422, detail="Tipo inválido. Opções: email, ip, domain.")

    value = req.value.lower().strip()
    if not value:
        raise HTTPException(status_code=422, detail="Valor não pode estar vazio.")

    try:
        sb = get_supabase_admin()
        result = sb.table("blacklist").insert({
            "type": req.type,
            "value": value,
            "reason": req.reason or f"Bloqueado por {admin_email}",
            "added_by": admin_email,
        }).execute()

        # Forçar recarga do cache
        with _blacklist_lock:
            _blacklist_cache["loaded_at"] = 0

        logger.info(f"[BLACKLIST] Adicionado: {req.type}={value} por {admin_email}")
        return {"status": "ok", "blocked": result.data[0] if result.data else {}}
    except Exception as e:
        logger.error(f"Erro ao adicionar à blacklist: {e}")
        raise HTTPException(status_code=500, detail="Erro ao adicionar à blacklist.")


@app.delete("/admin/blacklist/{entry_id}")
@limiter.limit("30/minute")
async def admin_blacklist_remove(
    entry_id: str,
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Remove entrada da blacklist (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(status_code=403, detail="Apenas administradores.")

    try:
        sb = get_supabase_admin()
        sb.table("blacklist").delete().eq("id", entry_id).execute()

        # Forçar recarga do cache
        with _blacklist_lock:
            _blacklist_cache["loaded_at"] = 0

        logger.info(f"[BLACKLIST] Removido: {entry_id} por {admin_email}")
        return {"status": "ok", "removed": entry_id}
    except Exception as e:
        logger.error(f"Erro ao remover da blacklist: {e}")
        raise HTTPException(status_code=500, detail="Erro ao remover da blacklist.")


# ============================================================
# ADMIN - ALERTAS DE SEGURANÇA
# ============================================================

@app.get("/admin/security-alerts")
@limiter.limit("30/minute")
async def admin_security_alerts(
    request: Request,
    resolved: bool = False,
    limit: int = 50,
    user: dict = Depends(get_current_user),
):
    """Lista alertas de segurança (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Apenas administradores.",
        )

    try:
        sb = get_supabase_admin()
        result = sb.table("security_alerts").select("*").eq(
            "resolved", resolved
        ).order("created_at", desc=True).limit(min(limit, 100)).execute()
        return {
            "alerts": result.data or [],
            "total": len(result.data or []),
            "showing": "resolved" if resolved else "unresolved",
        }
    except Exception as e:
        logger.error(f"Erro ao consultar alertas: {e}")
        raise HTTPException(status_code=500, detail="Erro ao consultar alertas.")


@app.post("/admin/security-alerts/{alert_id}/resolve")
@limiter.limit("30/minute")
async def resolve_security_alert(
    alert_id: str,
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Marca um alerta como resolvido (apenas admin)."""
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Apenas administradores.",
        )

    try:
        sb = get_supabase_admin()
        sb.table("security_alerts").update({
            "resolved": True,
        }).eq("id", alert_id).execute()
        return {"status": "ok", "alert_id": alert_id}
    except Exception as e:
        logger.error(f"Erro ao resolver alerta: {e}")
        raise HTTPException(status_code=500, detail="Erro ao resolver alerta.")


# ============================================================
# ADMIN - VERIFICAÇÃO DE PASSWORD
# ============================================================

class AdminVerifyRequest(BaseModel):
    password: str


@app.post("/admin/verify")
@limiter.limit("5/minute")
@limiter.limit("15/hour")
@limiter.limit("30/day")
async def admin_verify(request: Request, req: AdminVerifyRequest):
    """Verifica a password de admin para acesso ao painel de diagnóstico."""
    admin_password = os.environ.get("ADMIN_PASSWORD", "")

    if not admin_password:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Configuração de admin em falta.",
        )

    if not secrets.compare_digest(req.password, admin_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Password incorrecta.",
        )

    # Gerar token de sessão admin (válido 1 hora)
    admin_token = secrets.token_urlsafe(32)
    now = datetime.now(timezone.utc)
    with _admin_sessions_lock:
        _admin_sessions[admin_token] = {
            "created_at": now,
            "ip": get_remote_address(request),
        }
        # Cleanup expired sessions (older than 1 hour)
        expired = [k for k, v in _admin_sessions.items()
                   if (now - v["created_at"]).total_seconds() > ADMIN_SESSION_TTL]
        for k in expired:
            del _admin_sessions[k]
    return {"status": "ok", "message": "Acesso autorizado.", "admin_token": admin_token}


# ============================================================
# ADMIN - LOGS EM TEMPO REAL (monitorização remota)
# ============================================================

@app.get("/admin/logs")
@limiter.limit("60/minute")
async def admin_logs(
    request: Request,
    limit: int = 1000,
    level: Optional[str] = None,
    search: Optional[str] = None,
    since_id: int = 0,
    user: dict = Depends(get_current_user),
):
    """
    Retorna logs em memória para monitorização remota (apenas admin).

    Params:
        limit: máx entries (default 1000)
        level: filtrar por nível (INFO, WARNING, ERROR)
        search: filtrar por texto (case-insensitive)
        since_id: só logs com id > since_id (para polling incremental)
    """
    admin_email = (user.get("email", "") or "").lower().strip()
    if admin_email not in ADMIN_EMAILS:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Apenas administradores podem aceder aos logs.",
        )
    logs = _log_buffer.get_logs(limit=limit, level=level, search=search, since_id=since_id)
    return {
        "count": len(logs),
        "logs": logs,
    }

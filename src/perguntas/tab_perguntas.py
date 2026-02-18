# -*- coding: utf-8 -*-
"""
INTERFACE PERGUNTAS ADICIONAIS - VERSÃO ACUMULATIVA
═══════════════════════════════════════════════════════════════════════════

NOVO: Sistema ACUMULATIVO com upload de documentos!
- ✅ Upload múltiplos documentos (PDF, DOCX, TXT, XLSX)
- ✅ Extração automática de texto
- ✅ Histórico completo mantido
- ✅ Contexto NUNCA se perde
"""

import sys
from pathlib import Path

# Adicionar diretório raiz ao path (necessário para imports absolutos)
_ROOT = Path(__file__).resolve().parent.parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

import streamlit as st
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import logging
from io import BytesIO
import shutil

from src.utils.metadata_manager import listar_analises_com_titulos
from src.utils.sanitize import sanitize_run_id, sanitize_filename

# Imports para extração de documentos
try:
    import PyPDF2
    PDF_DISPONIVEL = True
except ImportError:
    PDF_DISPONIVEL = False
    logging.warning("PyPDF2 não disponível - PDFs não serão extraídos")

try:
    from docx import Document
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False
    logging.warning("python-docx não disponível - DOCX não serão extraídos")

try:
    import openpyxl
    XLSX_DISPONIVEL = True
except ImportError:
    XLSX_DISPONIVEL = False
    logging.warning("openpyxl não disponível - XLSX não serão extraídos")

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════
# FUNÇÕES EXTRAÇÃO DE TEXTO (NOVO!)
# ═══════════════════════════════════════════════════════════════════════════

def extrair_texto_pdf(file_bytes: bytes) -> str:
    """Extrai texto de PDF."""
    if not PDF_DISPONIVEL:
        return "[ERRO: PyPDF2 não instalado - não é possível extrair PDF]"
    
    try:
        pdf_file = BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        texto = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            if page_text.strip():
                texto.append(f"--- Página {page_num} ---\n{page_text}")
        
        return "\n\n".join(texto)
    
    except Exception as e:
        logger.error(f"Erro ao extrair PDF: {e}")
        return f"[ERRO AO EXTRAIR PDF: {e}]"


def extrair_texto_docx(file_bytes: bytes) -> str:
    """Extrai texto de DOCX."""
    if not DOCX_DISPONIVEL:
        return "[ERRO: python-docx não instalado - não é possível extrair DOCX]"
    
    try:
        docx_file = BytesIO(file_bytes)
        doc = Document(docx_file)
        
        texto = []
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                texto.append(para.text)
        
        return "\n\n".join(texto)
    
    except Exception as e:
        logger.error(f"Erro ao extrair DOCX: {e}")
        return f"[ERRO AO EXTRAIR DOCX: {e}]"


def extrair_texto_xlsx(file_bytes: bytes) -> str:
    """Extrai texto de XLSX."""
    if not XLSX_DISPONIVEL:
        return "[ERRO: openpyxl não instalado - não é possível extrair XLSX]"
    
    try:
        xlsx_file = BytesIO(file_bytes)
        wb = openpyxl.load_workbook(xlsx_file, data_only=True)
        
        texto = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            texto.append(f"=== FOLHA: {sheet_name} ===\n")
            
            for row in ws.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(row_values):
                    texto.append(" | ".join(row_values))
        
        return "\n".join(texto)
    
    except Exception as e:
        logger.error(f"Erro ao extrair XLSX: {e}")
        return f"[ERRO AO EXTRAIR XLSX: {e}]"


def extrair_texto_txt(file_bytes: bytes) -> str:
    """Extrai texto de TXT."""
    try:
        return file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return file_bytes.decode('latin-1')
        except Exception as e:
            logger.error(f"Erro ao extrair TXT: {e}")
            return f"[ERRO AO EXTRAIR TXT: {e}]"


def extrair_texto_documento(uploaded_file) -> str:
    """
    Extrai texto de documento anexado.
    
    Suporta: PDF, DOCX, XLSX, TXT
    
    Args:
        uploaded_file: streamlit.UploadedFile
    
    Returns:
        str: Texto extraído
    """
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name.lower()
    
    if file_name.endswith('.pdf'):
        return extrair_texto_pdf(file_bytes)
    elif file_name.endswith('.docx'):
        return extrair_texto_docx(file_bytes)
    elif file_name.endswith('.xlsx') or file_name.endswith('.xls'):
        return extrair_texto_xlsx(file_bytes)
    elif file_name.endswith('.txt'):
        return extrair_texto_txt(file_bytes)
    else:
        return f"[FORMATO NÃO SUPORTADO: {file_name}]"


# ═══════════════════════════════════════════════════════════════════════════
# FUNÇÕES AUXILIARES (mantidas do original)
# ═══════════════════════════════════════════════════════════════════════════

def detectar_ficheiros_soltos(output_dir: Path) -> bool:
    """Detecta se há ficheiros .md soltos diretos na pasta outputs."""
    if not output_dir.exists():
        return False
    ficheiros_md = list(output_dir.glob("*.md"))
    return len(ficheiros_md) > 0


def guardar_pergunta_resposta(
    run_id: str,
    output_dir: Path,
    pergunta: str,
    resultado,
    timestamp: str,
    documentos_anexados: List[str] = None  # ← NOVO!
) -> int:
    """
    Guarda pergunta e resposta PERMANENTEMENTE.
    
    ← MODIFICADO: Agora guarda resposta_final no JSON e documentos anexados!
    
    Args:
        run_id: ID da análise
        output_dir: Pasta outputs
        pergunta: Pergunta utilizador
        resultado: RespostaPergunta
        timestamp: Timestamp
        documentos_anexados: Lista nomes documentos anexados ← NOVO!
    
    Returns:
        int: Número da pergunta guardada
    """
    # Determinar pasta de destino
    if run_id == "__FICHEIROS_SOLTOS__":
        perguntas_dir = output_dir / "perguntas"
    else:
        run_id = sanitize_run_id(run_id)
        perguntas_dir = output_dir / run_id / "perguntas"

    perguntas_dir.mkdir(exist_ok=True, parents=True)

    # Contar perguntas existentes
    perguntas_existentes = list(perguntas_dir.glob("pergunta_*.json"))
    numero = len(perguntas_existentes) + 1
    
    # Nome base
    base_nome = f"pergunta_{numero}"
    
    # 1. ← MODIFICADO: Guardar metadata JSON (COM resposta_final e documentos!)
    metadata = {
        "numero": numero,
        "timestamp": timestamp,
        "pergunta": pergunta,
        "resposta_final": resultado.resposta_final,  # ← NOVO!
        "documentos_anexados": documentos_anexados or [],  # ← NOVO!
        "run_id": run_id,
        "tempo_ms": resultado.tempo_total_ms,
        "tokens": resultado.tokens_total,
        "custo": resultado.custo_estimado,
        "sucesso": resultado.sucesso
    }
    
    with open(perguntas_dir / f"{base_nome}.json", 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)
    
    # 2. Guardar COMPLETA (markdown)
    conteudo_completo = f"""# PERGUNTA #{numero}

**Data:** {timestamp}  
**Run ID:** {run_id}  
**Tempo:** {resultado.tempo_total_ms/1000:.1f}s  
**Tokens:** {resultado.tokens_total:,}  
**Custo:** ${resultado.custo_estimado:.4f}  
"""
    
    # ← NOVO: Adicionar documentos anexados
    if documentos_anexados and len(documentos_anexados) > 0:
        conteudo_completo += f"\n**Documentos Anexados:** {', '.join(documentos_anexados)}  \n"
    
    conteudo_completo += f"""
---

## 💭 PERGUNTA

{pergunta}

---

## 🔍 FASE 2: AUDITORIA CONSOLIDADA

{resultado.auditoria_consolidada}

---

## ⚖️ FASE 3: PARECERES JURÍDICOS

"""
    
    for i, juiz in enumerate(resultado.juizes, 1):
        conteudo_completo += f"""
### Relator {i} ({juiz.modelo})

{juiz.conteudo}

---
"""
    
    conteudo_completo += f"""
## 👨‍⚖️ FASE 4: PARECER FINAL DO CONSELHEIRO-MOR

{resultado.resposta_final}

---
"""
    
    with open(perguntas_dir / f"{base_nome}_completa.md", 'w', encoding='utf-8') as f:
        f.write(conteudo_completo)
    
    # 3. Guardar só auditoria
    with open(perguntas_dir / f"{base_nome}_auditoria.md", 'w', encoding='utf-8') as f:
        f.write(resultado.auditoria_consolidada)
    
    # 4. Guardar só decisão
    with open(perguntas_dir / f"{base_nome}_decisao.md", 'w', encoding='utf-8') as f:
        f.write(resultado.resposta_final)
    
    logger.info(f"✓ Pergunta #{numero} guardada em: {perguntas_dir}")
    
    return numero


def guardar_documentos_anexados(
    run_id: str,
    output_dir: Path,
    uploaded_files: List,
    textos_extraidos: Dict[str, str]
):
    """
    ← NOVA FUNÇÃO!
    
    Guarda documentos anexados PERMANENTEMENTE na pasta do projeto.
    
    Args:
        run_id: ID da análise
        output_dir: Pasta outputs
        uploaded_files: Lista UploadedFile do Streamlit
        textos_extraidos: Dict {nome_ficheiro: texto_extraido}
    """
    # Determinar pasta documentos
    if run_id == "__FICHEIROS_SOLTOS__":
        docs_dir = output_dir / "perguntas" / "documentos_anexados"
    else:
        run_id = sanitize_run_id(run_id)
        docs_dir = output_dir / run_id / "perguntas" / "documentos_anexados"

    docs_dir.mkdir(exist_ok=True, parents=True)

    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB limite por ficheiro

    for uploaded_file in uploaded_files:
        try:
            # Sanitizar nome do ficheiro para prevenir path traversal
            safe_name = sanitize_filename(uploaded_file.name)
            file_data = uploaded_file.getvalue()

            # Verificar tamanho do ficheiro
            if len(file_data) > MAX_FILE_SIZE:
                logger.warning(f"Ficheiro {safe_name} excede limite de {MAX_FILE_SIZE} bytes")
                continue

            # Guardar ficheiro original
            file_path = docs_dir / safe_name
            with open(file_path, 'wb') as f:
                f.write(file_data)

            # Guardar texto extraído
            texto = textos_extraidos.get(uploaded_file.name, "")
            nome_sem_ext = Path(safe_name).stem
            texto_path = docs_dir / f"{nome_sem_ext}_extraido.txt"
            
            with open(texto_path, 'w', encoding='utf-8') as f:
                f.write(texto)
            
            logger.info(f"Documento guardado: {safe_name}")

        except Exception as e:
            logger.error(f"Erro ao guardar documento: {e}")


# Continuação no próximo ficheiro devido ao tamanho...


# ═══════════════════════════════════════════════════════════════════════════
# FUNÇÕES EXPORTAÇÃO WORD/PDF (mantidas do original)
# ═══════════════════════════════════════════════════════════════════════════

def criar_word_auditoria(auditoria: str, pergunta: str) -> BytesIO:
    """Cria documento Word com auditoria."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('AUDITORIA CONSOLIDADA', 0)
        doc.add_paragraph(f'Pergunta: {pergunta}')
        doc.add_paragraph('')
        doc.add_paragraph(auditoria)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f'Erro ao criar Word: {e}')
        return None


def criar_word_decisao(decisao: str, pergunta: str) -> BytesIO:
    """Cria documento Word com decisão."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('DECISÃO FINAL', 0)
        doc.add_paragraph(f'Pergunta: {pergunta}')
        doc.add_paragraph('')
        doc.add_paragraph(decisao)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f'Erro ao criar Word: {e}')
        return None


# ═══════════════════════════════════════════════════════════════════════════
# PROCESSAMENTO (usa pipeline acumulativo!)
# ═══════════════════════════════════════════════════════════════════════════

def processar_pergunta_pipeline_completo(
    run_id: str,
    pergunta: str,
    output_dir: Path,
    auditor_models: List,
    juiz_models: List,
    presidente_model: str,
    llm_client,
    documentos_novos: List[Tuple[str, str]] = None,  # ← NOVO!
    chefe_model: str = None,  # v4.0 FIX: receber modelo como parâmetro
):
    """
    Processa pergunta usando pipeline completo.

    ← MODIFICADO: Agora aceita documentos_novos!
    """
    from src.perguntas.pipeline_perguntas import processar_pergunta_adicional

    return processar_pergunta_adicional(
        run_id=run_id,
        output_dir=output_dir,
        pergunta=pergunta,
        auditor_models=auditor_models,
        juiz_models=juiz_models,
        presidente_model=presidente_model,
        llm_client=llm_client,
        documentos_novos=documentos_novos,  # ← NOVO!
        chefe_model=chefe_model,  # v4.0 FIX: passar modelo explicitamente
    )


# ═══════════════════════════════════════════════════════════════════════════
# INTERFACE PRINCIPAL (MODIFICADA COM UPLOAD!)
# ═══════════════════════════════════════════════════════════════════════════

def tab_perguntas_adicionais(
    output_dir: Path,
    auditor_models: List,
    juiz_models: List,
    presidente_model: str,
    llm_client,
    chefe_model: str = None,
):
    """
    Interface principal - Perguntas Adicionais.
    
    ← MODIFICADO: Agora com upload de documentos!
    """
    st.title("💬 Perguntas Adicionais")
    
    st.markdown("""
    Faça perguntas sobre análises já processadas.  
    **NOVO:** Pode anexar documentos (minuta, comprovativo, etc.)!
    """)
    
    # ═════════════════════════════════════════════════════════════════════
    # SELECIONAR ANÁLISE EXISTENTE
    # ═════════════════════════════════════════════════════════════════════
    
    st.markdown("### 📂 Selecionar Análise")
    
    # ← NOVO: Usar função que retorna títulos
    analises_com_titulos = listar_analises_com_titulos(output_dir)
    
    if not analises_com_titulos:
        st.warning("⚠️ Nenhuma análise encontrada! Processe documentos primeiro.")
        return
    
    # Criar mapeamento: titulo_display -> run_id
    mapa_titulos = {}
    opcoes_display = []
    
    for run_id, titulo_display, data in analises_com_titulos:
        opcoes_display.append(f"📁 {titulo_display}")
        mapa_titulos[f"📁 {titulo_display}"] = run_id
    
    # Selectbox com títulos
    analise_selecionada_display = st.selectbox(
        "Escolha a análise:",
        opcoes_display,
        key="select_analise_perguntas"
    )
    
    # Obter run_id real
    run_id_selecionado = mapa_titulos[analise_selecionada_display]
    
    # ═════════════════════════════════════════════════════════════════════
    # ← NOVO: UPLOAD DOCUMENTOS ADICIONAIS (ACUMULATIVO)
    # ═════════════════════════════════════════════════════════════════════

    st.markdown("---")
    st.markdown("### 📎 Documentos Adicionais (Opcional)")

    st.info("💡 Adicione documentos um a um ou vários de cada vez. Os ficheiros são ACUMULADOS automaticamente.")

    # Inicializar session_state para ficheiros acumulados (perguntas)
    if "ficheiros_perguntas_acumulados" not in st.session_state:
        st.session_state.ficheiros_perguntas_acumulados = {}  # {nome: bytes}

    uploaded_files = st.file_uploader(
        "Adicionar documento(s):",
        accept_multiple_files=True,
        type=["pdf", "docx", "txt", "xlsx", "xls"],
        key="upload_docs_perguntas",
        help="Adicione ficheiros um a um ou vários. São acumulados automaticamente."
    )

    # Acumular novos ficheiros
    if uploaded_files:
        for f in uploaded_files:
            if f.name not in st.session_state.ficheiros_perguntas_acumulados:
                st.session_state.ficheiros_perguntas_acumulados[f.name] = f.getvalue()

    # Processar uploads acumulados
    documentos_extraidos = {}

    if st.session_state.ficheiros_perguntas_acumulados:
        col_titulo, col_limpar = st.columns([3, 1])
        with col_titulo:
            st.markdown(f"**{len(st.session_state.ficheiros_perguntas_acumulados)} documento(s) anexado(s):**")
        with col_limpar:
            if st.button("🗑️ Limpar", key="limpar_docs_perguntas"):
                st.session_state.ficheiros_perguntas_acumulados = {}
                st.rerun()

        ficheiros_a_remover = []
        for nome, dados in st.session_state.ficheiros_perguntas_acumulados.items():
            col1, col2, col3 = st.columns([3, 1, 0.5])

            with col1:
                st.write(f"📄 {nome} ({len(dados) / 1024:.1f} KB)")

            with col2:
                # Extrair texto do ficheiro
                from io import BytesIO
                class FicheiroPseudo:
                    def __init__(self, name, data):
                        self.name = name
                        self._data = data
                    def read(self):
                        return self._data

                texto = extrair_texto_documento(FicheiroPseudo(nome, dados))
                documentos_extraidos[nome] = texto

                if "[ERRO" in texto:
                    st.error("❌")
                else:
                    st.success(f"✅ {len(texto)} chars")

            with col3:
                if st.button("❌", key=f"rem_perg_{nome}", help=f"Remover {nome}"):
                    ficheiros_a_remover.append(nome)

        # Remover ficheiros marcados
        for nome in ficheiros_a_remover:
            del st.session_state.ficheiros_perguntas_acumulados[nome]
            st.rerun()

    # Criar lista de uploaded_files para compatibilidade com código existente
    uploaded_files = []
    if st.session_state.ficheiros_perguntas_acumulados:
        from io import BytesIO
        class FicheiroPseudoFull:
            def __init__(self, name, data):
                self.name = name
                self._data = data
                self.size = len(data)
            def read(self):
                return self._data
            def getvalue(self):
                return self._data

        uploaded_files = [
            FicheiroPseudoFull(nome, dados)
            for nome, dados in st.session_state.ficheiros_perguntas_acumulados.items()
        ]
    
    # ═════════════════════════════════════════════════════════════════════
    # HISTÓRICO PERGUNTAS (mostrar visualmente)
    # ═════════════════════════════════════════════════════════════════════
    
    st.markdown("---")
    st.markdown("### 📚 Histórico de Perguntas")
    
    # Determinar pasta perguntas
    if run_id_selecionado == "__FICHEIROS_SOLTOS__":
        perguntas_dir = output_dir / "perguntas"
    else:
        perguntas_dir = output_dir / run_id_selecionado / "perguntas"
    
    if perguntas_dir.exists():
        json_files = sorted(perguntas_dir.glob("pergunta_*.json"))
        
        if json_files:
            st.write(f"**{len(json_files)} pergunta(s) anterior(es):**")
            
            for json_file in json_files:
                try:
                    with open(json_file, 'r', encoding='utf-8') as f:
                        metadata = json.load(f)

                    with st.expander(f"❓ Pergunta #{metadata['numero']} ({metadata['timestamp']})"):
                        st.markdown(f"**Pergunta:** {metadata['pergunta']}")

                        if metadata.get('documentos_anexados'):
                            st.markdown(f"**Documentos:** {', '.join(metadata['documentos_anexados'])}")

                        # Tentar obter resposta do JSON ou do ficheiro .md
                        resposta = metadata.get('resposta_final')
                        if not resposta:
                            # Carregar do ficheiro _decisao.md (perguntas antigas)
                            decisao_file = json_file.parent / f"pergunta_{metadata['numero']}_decisao.md"
                            if decisao_file.exists():
                                with open(decisao_file, 'r', encoding='utf-8') as f:
                                    resposta = f.read()

                        st.markdown("**Resposta:**")
                        if resposta:
                            st.info(resposta[:2000] + "..." if len(resposta) > 2000 else resposta)
                        else:
                            st.warning("[Resposta não encontrada]")

                except Exception as e:
                    st.error(f"Erro ao carregar {json_file.name}: {e}")
        else:
            st.info("📝 Nenhuma pergunta anterior (primeira pergunta nesta análise)")
    else:
        st.info("📝 Nenhuma pergunta anterior (primeira pergunta nesta análise)")
    
    # ═════════════════════════════════════════════════════════════════════
    # NOVA PERGUNTA
    # ═════════════════════════════════════════════════════════════════════
    
    st.markdown("---")
    st.markdown("### ❓ Nova Pergunta")
    
    pergunta = st.text_area(
        "Escreva sua pergunta:",
        height=150,
        placeholder="Ex: Esta minuta de carta protege-me juridicamente? Devo alterar algo?",
        key="nova_pergunta_input"
    )
    
    # ═════════════════════════════════════════════════════════════════════
    # PROCESSAR
    # ═════════════════════════════════════════════════════════════════════
    
    if st.button("🚀 Processar Pergunta", type="primary", use_container_width=True):
        if not pergunta or len(pergunta.strip()) < 10:
            st.error("⚠️ Pergunta muito curta! Escreva pelo menos 10 caracteres.")
            return
        
        try:
            # Barra progresso (fake mas informativa)
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            status_text.text("🔄 Iniciando processamento...")
            progress_bar.progress(10)
            
            # ← NOVO: Preparar documentos para pipeline
            documentos_novos_lista = []
            if documentos_extraidos:
                for nome, texto in documentos_extraidos.items():
                    documentos_novos_lista.append((nome, texto))
            
            # ← NOVO: Guardar documentos ANTES de processar
            if uploaded_files and documentos_extraidos:
                status_text.text("💾 Guardando documentos anexados...")
                guardar_documentos_anexados(
                    run_id=run_id_selecionado,
                    output_dir=output_dir,
                    uploaded_files=uploaded_files,
                    textos_extraidos=documentos_extraidos
                )
                progress_bar.progress(20)
            
            status_text.text("🔍 Processando Fases 2-4 (pode demorar 3-5 min)...")
            progress_bar.progress(30)
            
            # PROCESSAR (COM DOCUMENTOS!)
            resultado = processar_pergunta_pipeline_completo(
                run_id=run_id_selecionado,
                pergunta=pergunta.strip(),
                output_dir=output_dir,
                auditor_models=auditor_models,
                juiz_models=juiz_models,
                presidente_model=presidente_model,
                llm_client=llm_client,
                documentos_novos=documentos_novos_lista,  # ← NOVO!
                chefe_model=chefe_model,  # v4.0 FIX: passar modelo explicitamente
            )
            
            progress_bar.progress(100)
            status_text.text("✅ Processamento concluído!")
            
            if not resultado.sucesso:
                st.error(f"❌ Erro: {resultado.erro}")
                return
            
            # ═════════════════════════════════════════════════════════════
            # GUARDAR PERMANENTEMENTE
            # ═════════════════════════════════════════════════════════════
            
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            numero_pergunta = guardar_pergunta_resposta(
                run_id=run_id_selecionado,
                output_dir=output_dir,
                pergunta=pergunta.strip(),
                resultado=resultado,
                timestamp=timestamp,
                documentos_anexados=[f.name for f in uploaded_files] if uploaded_files else []  # ← NOVO!
            )
            
            st.success(f"💾 Pergunta #{numero_pergunta} guardada permanentemente!")
            
            # ═════════════════════════════════════════════════════════════
            # MOSTRAR RESULTADOS
            # ═════════════════════════════════════════════════════════════
            
            st.divider()
            st.subheader("✅ Resposta do LexForum")
            
            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            
            with col_m1:
                st.metric("⏱️ Tempo", f"{resultado.tempo_total_ms/1000:.1f}s")
            with col_m2:
                st.metric("🔢 Tokens", f"{resultado.tokens_total:,}")
            with col_m3:
                st.metric("💰 Custo", f"${resultado.custo_estimado:.4f}")
            with col_m4:
                st.metric("🤖 LLMs", "7 (3+3+1)")
            
            st.markdown("---")
            
            st.markdown("### 👨‍⚖️ Parecer Final do Conselheiro-Mor")
            st.success(resultado.resposta_final)
            
            # ═════════════════════════════════════════════════════════════
            # BOTÕES DE EXPORTAÇÃO
            # ═════════════════════════════════════════════════════════════
            
            st.markdown("---")
            st.markdown("### 💾 Exportar Resultados")
            
            col_exp1, col_exp2, col_exp3, col_exp4 = st.columns(4)
            
            with col_exp1:
                buffer_aud_word = criar_word_auditoria(resultado.auditoria_consolidada, pergunta)
                if buffer_aud_word:
                    st.download_button(
                        label="📄 Auditoria (Word)",
                        data=buffer_aud_word,
                        file_name=f"auditoria_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            with col_exp2:
                buffer_dec_word = criar_word_decisao(resultado.resposta_final, pergunta)
                if buffer_dec_word:
                    st.download_button(
                        label="📄 Decisão (Word)",
                        data=buffer_dec_word,
                        file_name=f"decisao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            with col_exp3:
                st.download_button(
                    label="📝 Auditoria (TXT)",
                    data=resultado.auditoria_consolidada,
                    file_name=f"auditoria_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            with col_exp4:
                st.download_button(
                    label="📝 Decisão (TXT)",
                    data=resultado.resposta_final,
                    file_name=f"decisao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            st.markdown("---")
            
            with st.expander("📋 Ver Detalhes Completos"):
                st.markdown("#### Fase 2: Auditoria Consolidada")
                st.info(resultado.auditoria_consolidada)
                
                st.markdown("#### Fase 3: Pareceres Jurídicos")
                for i, juiz in enumerate(resultado.juizes, 1):
                    with st.expander(f"⚖️ Relator {i} ({juiz.modelo})"):
                        st.markdown(juiz.conteudo)
            
            st.success("✅ Pergunta processada, guardada e disponível para exportação!")
            st.info("💡 Pode rever esta resposta a qualquer momento no histórico acima!")
            
            # ← NOVO: Mostrar contexto acumulado
            if documentos_extraidos:
                st.info(f"📎 {len(documentos_extraidos)} documento(s) anexado(s) ao projeto e disponível(is) para futuras perguntas!")
            
        except Exception as e:
            st.error(f"❌ Erro: {e}")
            logger.error(f"Erro processando pergunta: {e}", exc_info=True)

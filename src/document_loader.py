# -*- coding: utf-8 -*-
"""
Carregador de documentos - Suporta PDF, DOCX, XLSX, TXT.
Extrai texto real dos ficheiros para análise.
"""

import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Union
from dataclasses import dataclass, field
from datetime import datetime
import hashlib
import io

from src.config import SUPPORTED_EXTENSIONS, LOG_LEVEL

# Configurar logging
logging.basicConfig(level=getattr(logging, LOG_LEVEL))
logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    """Conteúdo extraído de um documento."""
    filename: str
    extension: str
    text: str
    num_pages: int = 0
    num_chars: int = 0
    num_words: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_time: datetime = field(default_factory=datetime.now)
    file_hash: str = ""
    success: bool = True
    error: Optional[str] = None
    # PDF Seguro - campos adicionais
    pdf_safe_result: Optional[Any] = None  # PDFSafeResult quando aplicável
    pages_problematic: int = 0  # Contagem de páginas problemáticas
    pdf_safe_enabled: bool = False  # Se usou extração PDF Seguro

    def to_dict(self) -> Dict:
        return {
            "filename": self.filename,
            "extension": self.extension,
            "text": self.text[:1000] + "..." if len(self.text) > 1000 else self.text,
            "text_full_length": len(self.text),
            "num_pages": self.num_pages,
            "num_chars": self.num_chars,
            "num_words": self.num_words,
            "metadata": self.metadata,
            "extraction_time": self.extraction_time.isoformat(),
            "file_hash": self.file_hash,
            "success": self.success,
            "error": self.error,
            "pdf_safe_enabled": self.pdf_safe_enabled,
            "pages_problematic": self.pages_problematic,
        }


class DocumentLoader:
    """
    Carrega e extrai texto de documentos.

    Formatos suportados:
    - PDF (via pypdf)
    - DOCX (via python-docx)
    - XLSX (via openpyxl)
    - TXT (nativo)
    """

    def __init__(self):
        self._stats = {
            "total_loaded": 0,
            "successful": 0,
            "failed": 0,
            "by_extension": {},
        }

    def load(self, file_path: Union[str, Path, io.BytesIO], filename: Optional[str] = None) -> DocumentContent:
        """
        Carrega um documento e extrai o texto.

        Args:
            file_path: Caminho do ficheiro ou BytesIO
            filename: Nome do ficheiro (obrigatório se file_path for BytesIO)

        Returns:
            DocumentContent com o texto extraído
        """
        self._stats["total_loaded"] += 1

        # Determinar nome e extensão
        if isinstance(file_path, io.BytesIO):
            if not filename:
                raise ValueError("filename é obrigatório quando file_path é BytesIO")
            name = filename
            ext = Path(filename).suffix.lower()
            file_bytes = file_path.getvalue()
            file_hash = hashlib.md5(file_bytes).hexdigest()
        else:
            path = Path(file_path)
            name = path.name
            ext = path.suffix.lower()
            file_bytes = path.read_bytes()
            file_hash = hashlib.md5(file_bytes).hexdigest()

        # Verificar extensão suportada
        if ext not in SUPPORTED_EXTENSIONS:
            logger.error(f"Extensão não suportada: {ext}")
            self._stats["failed"] += 1
            return DocumentContent(
                filename=name,
                extension=ext,
                text="",
                success=False,
                error=f"Extensão não suportada: {ext}. Suportadas: {list(SUPPORTED_EXTENSIONS.keys())}",
            )

        # Atualizar estatísticas
        self._stats["by_extension"][ext] = self._stats["by_extension"].get(ext, 0) + 1

        try:
            # Extrair texto baseado na extensão
            if ext == ".pdf":
                text, pages, metadata = self._extract_pdf(file_bytes)
            elif ext == ".docx":
                text, pages, metadata = self._extract_docx(file_bytes)
            elif ext == ".xlsx":
                text, pages, metadata = self._extract_xlsx(file_bytes)
            elif ext == ".txt":
                text, pages, metadata = self._extract_txt(file_bytes)
            else:
                raise ValueError(f"Extrator não implementado para: {ext}")

            self._stats["successful"] += 1

            return DocumentContent(
                filename=name,
                extension=ext,
                text=text,
                num_pages=pages,
                num_chars=len(text),
                num_words=len(text.split()),
                metadata=metadata,
                file_hash=file_hash,
                success=True,
            )

        except Exception as e:
            logger.error(f"Erro ao extrair {name}: {e}")
            self._stats["failed"] += 1
            return DocumentContent(
                filename=name,
                extension=ext,
                text="",
                file_hash=file_hash,
                success=False,
                error=str(e),
            )

    def _extract_pdf(self, file_bytes: bytes) -> tuple:
        """Extrai texto de um PDF usando pdfplumber (melhor) ou pypdf (fallback)."""
        text = ""
        num_pages = 0
        metadata = {}

        # Tentar pdfplumber primeiro (melhor extração)
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                num_pages = len(pdf.pages)
                text_parts = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(f"[Página {i+1}]\n{page_text}")
                text = "\n\n".join(text_parts)
                metadata["extractor"] = "pdfplumber"
                logger.info(f"PDF extraído com pdfplumber: {num_pages} páginas, {len(text)} caracteres")
        except Exception as e:
            logger.warning(f"pdfplumber falhou: {e}, tentando pypdf...")
            # Fallback para pypdf
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                num_pages = len(reader.pages)
                text_parts = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(f"[Página {i+1}]\n{page_text}")
                text = "\n\n".join(text_parts)
                metadata["extractor"] = "pypdf"
                if reader.metadata:
                    for key in ["/Title", "/Author", "/Subject", "/Creator", "/Producer"]:
                        if key in reader.metadata:
                            metadata[key.replace("/", "")] = reader.metadata[key]
                logger.info(f"PDF extraído com pypdf: {num_pages} páginas, {len(text)} caracteres")
            except ImportError:
                raise ImportError("Nenhum extrator PDF disponível. Execute: pip install pdfplumber pypdf")

        # AVISO CRÍTICO: PDF sem texto
        if not text.strip():
            logger.warning(f"⚠️ PDF tem {num_pages} páginas mas 0 caracteres! Provavelmente é imagem escaneada.")
            metadata["aviso"] = "PDF sem texto extraível - possível imagem escaneada"

        return text, num_pages, metadata

    def _extract_docx(self, file_bytes: bytes) -> tuple:
        """Extrai texto de um DOCX usando python-docx com detecção de páginas."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx não instalado. Execute: pip install python-docx")

        doc = Document(io.BytesIO(file_bytes))

        # Detectar page breaks explícitos no XML do DOCX
        from lxml import etree
        BREAK_TYPE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'
        BREAK_TAG = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'
        LAST_RENDERED_PAGE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lastRenderedPageBreak'

        page_num = 1
        text_parts = []
        text_parts.append(f"[Página {page_num}]")

        # Extrair parágrafos com detecção de page breaks
        for para in doc.paragraphs:
            # Verificar page breaks no XML do parágrafo
            for run in para.runs:
                for elem in run._element:
                    if elem.tag == BREAK_TAG and elem.get(BREAK_TYPE) == 'page':
                        page_num += 1
                        text_parts.append(f"\n[Página {page_num}]")
                    elif elem.tag == LAST_RENDERED_PAGE:
                        page_num += 1
                        text_parts.append(f"\n[Página {page_num}]")
            if para.text.strip():
                text_parts.append(para.text)

        # Se não detectou nenhum page break, inserir marcadores sintéticos
        if page_num == 1:
            SYNTHETIC_PAGE_CHARS = 3000
            raw_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    raw_parts.append(para.text)
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    raw_parts.append("[TABELA]\n" + "\n".join(table_text))

            full_text = "\n\n".join(raw_parts)
            page_num = 1
            text_parts = [f"[Página {page_num}]"]
            char_count = 0
            for part in raw_parts:
                char_count += len(part) + 2  # +2 para \n\n
                if char_count > SYNTHETIC_PAGE_CHARS * page_num:
                    page_num += 1
                    text_parts.append(f"\n[Página {page_num}]")
                text_parts.append(part)
        else:
            # Tem page breaks reais — adicionar tabelas ao final
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("[TABELA]\n" + "\n".join(table_text))

        # Fix: se poucas páginas foram detectadas mas o texto é muito longo,
        # re-paginar com marcadores sintéticos (ex: DOCX com 184K chars e 2 páginas)
        total_chars = sum(len(p) for p in text_parts)
        avg_chars_per_page = total_chars / max(page_num, 1)
        SYNTHETIC_PAGE_CHARS = 3000
        if avg_chars_per_page > SYNTHETIC_PAGE_CHARS * 3:  # >9K por página = precisa subdividir
            logger.info(
                f"DOCX: {page_num} páginas detectadas mas {total_chars} chars "
                f"({avg_chars_per_page:.0f} chars/pág). Re-paginando sinteticamente."
            )
            raw_text = "\n\n".join(p for p in text_parts if not p.strip().startswith("[Página"))
            page_num = 1
            text_parts = [f"[Página {page_num}]"]
            char_count = 0
            for line in raw_text.split("\n\n"):
                if not line.strip():
                    continue
                char_count += len(line) + 2
                if char_count > SYNTHETIC_PAGE_CHARS * page_num:
                    page_num += 1
                    text_parts.append(f"\n[Página {page_num}]")
                text_parts.append(line)

        text = "\n\n".join(text_parts)

        # Metadata
        metadata = {}
        core_props = doc.core_properties
        if core_props.title:
            metadata["Title"] = core_props.title
        if core_props.author:
            metadata["Author"] = core_props.author
        if core_props.subject:
            metadata["Subject"] = core_props.subject
        if core_props.created:
            metadata["Created"] = core_props.created.isoformat() if core_props.created else None

        logger.info(f"DOCX extraído: {len(doc.paragraphs)} parágrafos, {page_num} páginas, {len(text)} caracteres")
        return text, page_num, metadata

    def _extract_xlsx(self, file_bytes: bytes) -> tuple:
        """Extrai texto de um XLSX usando openpyxl."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("openpyxl não instalado. Execute: pip install openpyxl")

        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)

        text_parts = []
        num_sheets = len(wb.sheetnames)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_text = [f"[FOLHA: {sheet_name}]"]

            for row in ws.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    sheet_text.append(" | ".join(row_values))

            if len(sheet_text) > 1:
                text_parts.append("\n".join(sheet_text))

        text = "\n\n".join(text_parts)

        metadata = {
            "num_sheets": num_sheets,
            "sheet_names": wb.sheetnames,
        }

        logger.info(f"XLSX extraído: {num_sheets} folhas, {len(text)} caracteres")
        return text, num_sheets, metadata

    def _extract_txt(self, file_bytes: bytes) -> tuple:
        """Extrai texto de um ficheiro TXT."""
        # Tentar diferentes encodings
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

        text = None
        used_encoding = None

        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            # Fallback: ignorar erros
            text = file_bytes.decode("utf-8", errors="ignore")
            used_encoding = "utf-8 (com erros)"

        # Contar linhas como "páginas"
        num_lines = text.count("\n") + 1

        metadata = {
            "encoding": used_encoding,
            "num_lines": num_lines,
        }

        logger.info(f"TXT extraído: {num_lines} linhas, {len(text)} caracteres")
        return text, 1, metadata

    def load_multiple(self, file_paths: List[Union[str, Path]]) -> List[DocumentContent]:
        """Carrega múltiplos documentos."""
        return [self.load(fp) for fp in file_paths]

    def load_pdf_safe(
        self,
        file_path: Union[str, Path, io.BytesIO],
        filename: Optional[str] = None,
        out_dir: Optional[Path] = None
    ) -> DocumentContent:
        """
        Carrega PDF usando o sistema PDF Seguro (página-a-página).

        Args:
            file_path: Caminho do ficheiro ou BytesIO
            filename: Nome do ficheiro (obrigatório se file_path for BytesIO)
            out_dir: Diretório para outputs (páginas, manifest)

        Returns:
            DocumentContent com pdf_safe_result preenchido
        """
        from src.pipeline.pdf_safe import get_pdf_safe_loader, PDFSafeResult

        self._stats["total_loaded"] += 1

        # Determinar nome e bytes
        if isinstance(file_path, io.BytesIO):
            if not filename:
                raise ValueError("filename é obrigatório quando file_path é BytesIO")
            name = filename
            file_bytes = file_path.getvalue()
        else:
            path = Path(file_path)
            name = path.name
            file_bytes = path.read_bytes()

        file_hash = hashlib.md5(file_bytes).hexdigest()
        ext = Path(name).suffix.lower()

        if ext != ".pdf":
            logger.warning(f"PDF Seguro só suporta PDFs, usando loader normal para {ext}")
            return self.load(file_path, filename)

        # Criar diretório de output se não especificado
        if out_dir is None:
            from src.config import OUTPUT_DIR
            import uuid
            out_dir = OUTPUT_DIR / f"pdf_safe_{uuid.uuid4().hex[:8]}"

        out_dir.mkdir(parents=True, exist_ok=True)

        try:
            # Usar PDF Seguro (com Vision OCR se LLM client disponível)
            try:
                from src.llm_client import get_llm_client
                llm_client = get_llm_client()
            except Exception:
                llm_client = None
            loader = get_pdf_safe_loader(llm_client=llm_client)
            pdf_result = loader.load_pdf_pages(file_bytes, name, out_dir)

            # Combinar texto de todas as páginas
            text_parts = []
            scanned_pages = {}  # page_num → image_path para páginas escaneadas
            for page in pdf_result.pages:
                # Usar override se existir, senão text_clean
                page_text = page.override_text if page.override_text else page.text_clean
                if page_text.strip():
                    text_parts.append(f"[Página {page.page_num}]\n{page_text}")
                # Registar páginas escaneadas (VISUAL_PENDING) para análise visual
                if (page.status_final == "VISUAL_PENDING"
                        and page.image_path
                        and Path(page.image_path).exists()):
                    scanned_pages[page.page_num] = str(page.image_path)

            text = "\n\n".join(text_parts)

            if scanned_pages:
                logger.info(
                    f"📸 {len(scanned_pages)} página(s) escaneada(s) detetada(s) em {name}: "
                    f"páginas {list(scanned_pages.keys())}"
                )

            self._stats["successful"] += 1
            self._stats["by_extension"][ext] = self._stats["by_extension"].get(ext, 0) + 1

            return DocumentContent(
                filename=name,
                extension=ext,
                text=text,
                num_pages=pdf_result.total_pages,
                num_chars=len(text),
                num_words=len(text.split()),
                metadata={
                    "extractor": "pdf_safe",
                    "pages_ok": pdf_result.pages_ok,
                    "pages_suspeita": pdf_result.pages_suspeita,
                    "pages_sem_texto": pdf_result.pages_sem_texto,
                    "document_provenance": pdf_result.document_provenance,
                    "scanned_pages": scanned_pages,
                },
                file_hash=file_hash,
                success=True,
                pdf_safe_result=pdf_result,
                pages_problematic=pdf_result.pages_suspeita + pdf_result.pages_sem_texto,
                pdf_safe_enabled=True,
            )

        except Exception as e:
            logger.error(f"Erro no PDF Seguro para {name}: {e}")
            self._stats["failed"] += 1
            return DocumentContent(
                filename=name,
                extension=ext,
                text="",
                file_hash=file_hash,
                success=False,
                error=str(e),
            )

    def get_stats(self) -> Dict:
        """Retorna estatísticas de carregamento."""
        return self._stats.copy()

    def reset_stats(self):
        """Reseta as estatísticas."""
        self._stats = {
            "total_loaded": 0,
            "successful": 0,
            "failed": 0,
            "by_extension": {},
        }


# Instância global
_global_loader: Optional[DocumentLoader] = None


def get_document_loader() -> DocumentLoader:
    """Retorna o carregador de documentos global."""
    global _global_loader
    if _global_loader is None:
        _global_loader = DocumentLoader()
    return _global_loader


def load_document(file_path: Union[str, Path, io.BytesIO], filename: Optional[str] = None) -> DocumentContent:
    """Função de conveniência para carregar um documento."""
    loader = get_document_loader()
    return loader.load(file_path, filename)


def get_supported_extensions() -> Dict[str, str]:
    """Retorna as extensões suportadas."""
    return SUPPORTED_EXTENSIONS.copy()
